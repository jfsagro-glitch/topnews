"""
Основной Telegram бот для публикации новостей
"""
import logging
import time
import os
import tempfile
import socket
import hmac
import hashlib
import secrets
import json
from contextlib import suppress
from datetime import datetime
from net.deepseek_client import DeepSeekClient
from urllib.parse import urlparse
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, CallbackQueryHandler, MessageHandler, filters,
    ContextTypes, ConversationHandler
)
from telegram.constants import ParseMode
from telegram.error import Conflict, RetryAfter, TimedOut
import asyncio
from config.config import TELEGRAM_TOKEN, TELEGRAM_CHANNEL_ID, CHECK_INTERVAL_SECONDS, ADMIN_IDS, AI_CALLS_PER_TICK_MAX
from utils.env import get_app_env

logger = logging.getLogger(__name__)

# Import DATABASE_PATH/ACCESS_DB_PATH from railway_config if available, else from config
try:
    from config.railway_config import DATABASE_PATH, ACCESS_DB_PATH, INVITE_SECRET
except (ImportError, ValueError):
    from config.config import DATABASE_PATH, ACCESS_DB_PATH, INVITE_SECRET

try:
    from config.railway_config import SOURCES_CONFIG as ACTIVE_SOURCES_CONFIG
except (ImportError, ValueError):
    from config.config import SOURCES_CONFIG as ACTIVE_SOURCES_CONFIG

from db.database import NewsDatabase
from utils.text_cleaner import format_telegram_message
from utils.content_quality import compute_simhash, compute_url_hash, hamming_distance, normalize_url
from utils.date_parser import get_project_now, parse_datetime_value, parse_url_date, to_project_tz
from sources.source_collector import SourceCollector
from core.services.access_control import AILevelManager, get_llm_profile
from core.services.ai_gate import AITickGate
from core.services.collection_stop import (
    get_global_collection_stop_state,
    get_global_collection_stop_status,
    set_global_collection_stop,
)
from utils.mgmt_api import start_mgmt_api, stop_mgmt_api


class NewsBot:
    """Основной класс Telegram бота"""
    
    def __init__(self):
        self.application = None
        self.db = NewsDatabase(db_path=DATABASE_PATH)  # News DB (env-specific)
        self.access_db = NewsDatabase(db_path=ACCESS_DB_PATH)  # Shared access control DB
        
        # DeepSeek client with cache and budget enabled
        from config.config import DEEPSEEK_API_KEY
        self.deepseek_client = DeepSeekClient(api_key=DEEPSEEK_API_KEY, db=self.db)
        
        # AI category verification toggle (can be controlled via button)
        from config.config import AI_CATEGORY_VERIFICATION_ENABLED
        self.ai_verification_enabled = AI_CATEGORY_VERIFICATION_ENABLED
        
        # SourceCollector with optional AI verification
        self.collector = SourceCollector(db=self.db, ai_client=self.deepseek_client, bot=self)
        
        # Initialize sources from SOURCES_CONFIG
        self._init_sources()
        
        self.is_running = True
        self.is_paused = False
        self.collection_lock = asyncio.Lock()  # Prevent concurrent collection cycles
        
        # Cache for recently published news (for AI button)
        self.news_cache = {}  # news_id -> {'title', 'text', 'source', 'url'}
        
        # Admin ids (from environment or config fallback)
        self.admin_ids = self._load_admin_ids()
        
        # Rate limiting for AI summarize requests (per user per minute)
        self.user_ai_requests = {}  # {user_id: [timestamp1, timestamp2, ...]}

        # Per-tick AI gating
        self._ai_tick_gate = AITickGate(max_calls=AI_CALLS_PER_TICK_MAX)
        self._ai_tick_id = None

        # Drop reasons counters (domain -> reason -> count)
        self.drop_counters = {}

        # Track last "no news" notification per user to avoid duplicates in short bursts
        self._last_no_news_sent_at = {}
        
        # Instance lock (prevent double start)
        self._instance_lock_fd = None
        self._instance_lock_path = None
        self._db_instance_id = f"{socket.gethostname()}:{os.getpid()}"
        self._shutdown_requested = False

    def _is_admin(self, user_id: int) -> bool:
        """Check if user is admin (hardcoded ADMIN_IDS or config ADMIN_USER_IDS)."""
        admin_ids = set(self.admin_ids)
        
        # Add admins from config
        try:
            from config.railway_config import ADMIN_USER_IDS
        except (ImportError, ValueError):
            from config.config import ADMIN_USER_IDS
        if ADMIN_USER_IDS:
            admin_ids.update(ADMIN_USER_IDS)
        
        return user_id in admin_ids

    def _has_access(self, user_id: int) -> bool:
        """Check if user has access to bot (admin or approved via invite)."""
        app_env = get_app_env()
        
        # Admins always have access
        if self._is_admin(user_id):
            return True
        
        # Sandbox is open to all
        if app_env == "sandbox":
            return self._is_admin(user_id)
        
        # Prod requires approval via invite
        return self.access_db.is_user_approved(str(user_id))

    def _check_access(self, handler):
        """Decorator to check user access before executing handler"""
        async def wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE):
            user_id = update.effective_user.id

            # Sandbox is admin-only
            if get_app_env() == "sandbox" and not self._is_admin(user_id):
                if update.message:
                    await update.message.reply_text("⛔ Access denied")
                return
            
            # /start can always be called (handles invite codes and access messages)
            if handler.__name__ == 'cmd_start':
                return await handler(update, context)
            
            # Check if user has access
            if not self._has_access(user_id):
                try:
                    from config.railway_config import APP_ENV
                except (ImportError, ValueError):
                    from config.config import APP_ENV
                
                await update.message.reply_text(
                    "🔒 Доступ к боту только по инвайту.\n\n"
                    "Для получения доступа:\n"
                    "1. Обратитесь к администратору\n"
                    "2. Получите инвайт-ссылку\n"
                    "3. Перейдите по ссылке для активации"
                )
                return
            
            # User has access, execute handler
            return await handler(update, context)
        
        return wrapper

    def _load_admin_ids(self) -> list[int]:
        raw_list = os.getenv("ADMIN_TELEGRAM_IDS", "") or ""
        raw_single = os.getenv("ADMIN_TELEGRAM_ID", "") or ""
        values = []
        for item in raw_list.split(","):
            item = item.strip()
            if item.isdigit():
                values.append(int(item))
        if raw_single.strip().isdigit():
            values.append(int(raw_single.strip()))
        if values:
            return sorted(set(values))
        return list(ADMIN_IDS)

    async def _sandbox_admin_guard(self, update: Update | None = None, query=None) -> bool:
        """Block non-admin access in sandbox for any command/callback."""
        if get_app_env() != "sandbox":
            return True
        user_id = None
        if query is not None:
            user_id = query.from_user.id
        elif update is not None and update.effective_user:
            user_id = update.effective_user.id
        if user_id is not None and self._is_admin(user_id):
            return True
        if query is not None:
            await query.answer("⛔ Access denied", show_alert=True)
        elif update is not None and update.message:
            await update.message.reply_text("⛔ Access denied")
        return False

    def _init_admins_access(self):
        """Initialize admin users with access to prod bot"""
        for admin_id in self.admin_ids:
            # Check if already approved
            if not self.access_db.is_user_approved(str(admin_id)):
                # Add admin with "SYSTEM" as invited_by
                from datetime import datetime
                cursor = self.access_db._conn.cursor()
                with self.access_db._write_lock:
                    cursor.execute(
                        'INSERT OR IGNORE INTO approved_users (user_id, username, first_name, invited_by, approved_at) VALUES (?, ?, ?, ?, ?)',
                        (str(admin_id), None, f"Admin {admin_id}", "SYSTEM", datetime.now().isoformat())
                    )
                    self.access_db._conn.commit()
                logger.info(f"Initialized admin access for user {admin_id}")

    def _get_sandbox_filter_user_id(self) -> int | None:
        """Pick a user id whose source settings control sandbox filtering."""
        return self.admin_ids[0] if self.admin_ids else None

    def _begin_ai_tick(self, tick_id: str) -> None:
        if self._ai_tick_gate:
            self._ai_tick_gate.begin_tick(tick_id)
        self._ai_tick_id = tick_id

    def _ai_tick_allow(self, task: str) -> bool:
        if not self._ai_tick_gate:
            return True
        if not self._ai_tick_gate.can_call(task):
            return False
        self._ai_tick_gate.record_call(task)
        return True

    def _get_ai_tick_state(self) -> dict:
        if not self._ai_tick_gate:
            return {"tick_id": None, "calls": 0, "max_calls": 0, "disabled": []}
        return self._ai_tick_gate.get_state()
    
    def _init_sources(self):
        """Инициализировать список источников из ACTIVE_SOURCES_CONFIG"""
        try:
            if not hasattr(self.db, "get_or_create_sources"):
                logger.warning("Source initialization skipped: get_or_create_sources not available")
                return
            sources_to_create = []
            
            # Собрать все источники из конфига, обрабатывая ВСЕ категории одинаково
            for category, cfg in ACTIVE_SOURCES_CONFIG.items():
                for src_url in cfg.get('sources', []):
                    # Telegram каналы - используем имя канала как код
                    if 't.me' in src_url:
                        channel = src_url.replace('https://t.me/', '').replace('http://t.me/', '').replace('@', '').strip('/')
                        if channel:
                            sources_to_create.append({'code': channel, 'title': f"@{channel}"})
                    else:
                        # Web источники (по домену)
                        domain = src_url.replace('https://', '').replace('http://', '').split('/')[0]
                        if domain:
                            sources_to_create.append({'code': domain, 'title': domain})
            
            # Убрать дубликаты
            seen_codes = set()
            unique_sources = []
            for src in sources_to_create:
                if src['code'] not in seen_codes:
                    unique_sources.append(src)
                    seen_codes.add(src['code'])
            
            # Создать или обновить в БД
            self.db.get_or_create_sources(unique_sources)
            logger.info(f"Initialized {len(unique_sources)} sources in database")
        except Exception as e:
            logger.error(f"Error initializing sources: {e}")

    def _acquire_instance_lock(self) -> bool:
        """Acquire a filesystem lock to prevent multiple bot instances."""
        try:
            lock_dir = tempfile.gettempdir()
            lock_path = os.path.join(lock_dir, "topnews_bot.lock")
            self._instance_lock_path = lock_path

            # In sandbox, always clear stale lock to avoid restart loops
            try:
                from config.config import APP_ENV
                if APP_ENV == "sandbox" and os.path.exists(lock_path):
                    os.remove(lock_path)
            except Exception:
                pass

            # If lock exists, attempt to validate PID liveness
            if os.path.exists(lock_path):
                try:
                    with open(lock_path, "r", encoding="utf-8") as fh:
                        raw_pid = fh.read().strip()
                    pid = int(raw_pid) if raw_pid.isdigit() else None
                except Exception:
                    pid = None

                def _pid_is_running(check_pid: int | None) -> bool:
                    if not check_pid or check_pid <= 0:
                        return False
                    try:
                        os.kill(check_pid, 0)
                    except Exception:
                        return False
                    return True

                if pid and not _pid_is_running(pid):
                    logger.warning("Stale instance lock with dead PID found. Removing.")
                    try:
                        os.remove(lock_path)
                    except Exception:
                        pass
                elif pid and _pid_is_running(pid):
                    logger.error("Another bot instance appears to be running. Exiting.")
                    return False

            # If stale lock older than 30 minutes, remove it as a fallback
            stale_seconds = 30 * 60
            if os.path.exists(lock_path):
                try:
                    mtime = os.path.getmtime(lock_path)
                    if time.time() - mtime > stale_seconds:
                        logger.warning("Stale instance lock found (timeout). Removing.")
                        os.remove(lock_path)
                except Exception:
                    pass

            fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_RDWR)
            self._instance_lock_fd = fd
            os.write(fd, str(os.getpid()).encode("utf-8"))
            return True
        except FileExistsError:
            logger.error("Another bot instance appears to be running. Exiting.")
            return False
        except Exception as e:
            logger.error(f"Failed to acquire instance lock: {e}")
            return False

    def _get_invite_secret(self) -> str | None:
        """Read invite secret dynamically (supports env updates without code changes)."""
        env_val = os.getenv('INVITE_SECRET')
        if env_val and env_val.strip():
            return env_val.strip()
        if INVITE_SECRET and str(INVITE_SECRET).strip():
            return str(INVITE_SECRET).strip()
        return None

    def _generate_signed_invite_code(self, created_by: str) -> str | None:
        """Generate a signed invite code that can be verified without shared DB."""
        secret = self._get_invite_secret()
        if not secret:
            logger.error("INVITE_SECRET not set; cannot generate signed invite")
            return None
        try:
            payload = secrets.token_urlsafe(8)
            sig = hmac.new(secret.encode(), payload.encode(), hashlib.sha256).hexdigest()[:10]
            return f"{payload}-{sig}"
        except Exception as e:
            logger.error(f"Error generating signed invite: {e}")
            return None

    def _release_instance_lock(self):
        """Release filesystem instance lock."""
        try:
            if self._instance_lock_fd is not None:
                os.close(self._instance_lock_fd)
                self._instance_lock_fd = None
            if self._instance_lock_path and os.path.exists(self._instance_lock_path):
                os.remove(self._instance_lock_path)
        except Exception as e:
            logger.debug(f"Failed to release instance lock: {e}")

    def create_application(self) -> Application:
        """Создает и конфигурирует Telegram Application"""
        
        self.application = Application.builder().token(TELEGRAM_TOKEN).build()
        
        # Регистрируем обработчики команд с проверкой доступа
        self.application.add_handler(CommandHandler("start", self._check_access(self.cmd_start)))
        self.application.add_handler(CommandHandler("help", self._check_access(self.cmd_help)))
        self.application.add_handler(CommandHandler("sync", self._check_access(self.cmd_sync)))
        self.application.add_handler(CommandHandler("status", self._check_access(self.cmd_status)))
        self.application.add_handler(CommandHandler("pause", self._check_access(self.cmd_pause)))
        self.application.add_handler(CommandHandler("resume", self._check_access(self.cmd_resume)))
        self.application.add_handler(CommandHandler("filter", self._check_access(self.cmd_filter)))
        self.application.add_handler(CommandHandler("sync_deepseek", self._check_access(self.cmd_sync_deepseek)))
        self.application.add_handler(CommandHandler("update_stats", self._check_access(self.cmd_update_stats)))
        self.application.add_handler(CommandHandler("debug_sources", self._check_access(self.cmd_debug_sources)))
        self.application.add_handler(CommandHandler("my_selection", self._check_access(self.cmd_my_selection)))
        
        # Обработчик текстовых сообщений (эмодзи-кнопки)
        self.application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, self.handle_emoji_buttons))
        
        # Обработчик inline кнопок
        self.application.add_handler(CallbackQueryHandler(self.button_callback))

        # Global error handler
        self.application.add_error_handler(self.on_error)
        
        logger.info("Application created successfully")
        return self.application

    # Persistent reply keyboard for chats (anchored at bottom)
    # For regular users (prod)
    REPLY_KEYBOARD = ReplyKeyboardMarkup(
        [['🔄', '✉️', '⏸️', '▶️'], ['⚙️ Настройки']], resize_keyboard=True, one_time_keyboard=False
    )

    def _build_sandbox_admin_keyboard(self) -> InlineKeyboardMarkup:
        from core.services.global_stop import get_global_stop

        is_stopped = get_global_stop()
        
        # Если остановлена - показываем "Возобновить работу", иначе - "Остановить систему"
        if is_stopped:
            keyboard = [
                [InlineKeyboardButton("▶️ Возобновить работу", callback_data="mgmt:resume_work")],
            ]
        else:
            keyboard = [
                [InlineKeyboardButton("⛔ ОСТАНОВИТЬ ВСЮ СИСТЕМУ", callback_data="mgmt:toggle_global_stop")],
            ]
        
        keyboard.extend([
            [InlineKeyboardButton("📊 Статус системы", callback_data="mgmt:status")],
            [InlineKeyboardButton("🤖 AI управление", callback_data="mgmt:ai")],
            [InlineKeyboardButton("📰 Источники", callback_data="mgmt:sources")],
            [InlineKeyboardButton("📈 Статистика", callback_data="mgmt:stats")],
            [InlineKeyboardButton("⚙ Настройки", callback_data="mgmt:settings")],
            [InlineKeyboardButton("👥 Пользователи и инвайты", callback_data="mgmt:users")],
            [InlineKeyboardButton("🧰 Диагностика", callback_data="mgmt:diag")],
            [InlineKeyboardButton("↩️ Назад", callback_data="mgmt:main")],
        ])

        return InlineKeyboardMarkup(keyboard)

    def _get_rsshub_telegram_enabled(self) -> bool:
        try:
            value = self.db.get_system_setting("rsshub_telegram_enabled")
        except Exception:
            value = None
        if value is None:
            try:
                from config.railway_config import RSSHUB_TELEGRAM_ENABLED
            except (ImportError, ValueError):
                from config.config import RSSHUB_TELEGRAM_ENABLED
            return bool(RSSHUB_TELEGRAM_ENABLED)
        return value.strip() not in ("0", "false", "False")
    
    async def cmd_start(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /start"""
        try:
            from config.railway_config import APP_ENV
        except (ImportError, ValueError):
            from config.config import APP_ENV
        
        user_id = update.message.from_user.id
        username = update.message.from_user.username
        first_name = update.message.from_user.first_name
        
        # Проверка инвайт-кода (если передан через deep link)
        if context.args and len(context.args) > 0:
            invite_code = context.args[0]
            
            # Попытка использовать инвайт через access БД
            if self.access_db.use_invite(invite_code, str(user_id), username, first_name):
                await update.message.reply_text(
                    "✅ Инвайт-код успешно активирован!\n\n"
                    "Теперь у вас есть доступ к боту. Используйте /help для списка команд.",
                    reply_markup=ReplyKeyboardRemove() if APP_ENV == "sandbox" else self.REPLY_KEYBOARD
                )
                return

            # Если это подписанный инвайт, проверяем подпись
            secret = self._get_invite_secret()
            if '-' in invite_code and not secret:
                await update.message.reply_text(
                    "❌ Не задан INVITE_SECRET в проде.\n\n"
                    "Инвайт создан как подписанный, но секрет отсутствует.\n"
                    "Установите INVITE_SECRET одинаково в prod и sandbox и перезапустите бота."
                )
                return

            if '-' in invite_code and secret:
                if self.access_db.use_signed_invite(invite_code, str(user_id), username, first_name, secret):
                    await update.message.reply_text(
                        "✅ Инвайт-код успешно активирован!\n\n"
                        "Теперь у вас есть доступ к боту. Используйте /help для списка команд.",
                        reply_markup=ReplyKeyboardRemove() if APP_ENV == "sandbox" else self.REPLY_KEYBOARD
                    )
                    return
                else:
                    await update.message.reply_text(
                        "❌ Инвайт-код неверный или подпись не совпадает.\n\n"
                        "Проверьте, что инвайт создан в песочнице после обновления и что INVITE_SECRET одинаковый в prod и sandbox."
                    )
                    return
            else:
                await update.message.reply_text(
                    "❌ Неверный или уже использованный инвайт-код.\n\n"
                    "Получите новый инвайт от администратора."
                )
                return
        
        # Проверка доступа
        if not self._has_access(user_id):
            await update.message.reply_text(
                "🔒 Доступ к боту только по инвайту.\n\n"
                "Для получения доступа:\n"
                "1. Обратитесь к администратору\n"
                "2. Получите инвайт-ссылку\n"
                "3. Перейдите по ссылке для активации"
            )
            return
        
        is_admin = self._is_admin(user_id)
        env_marker = "\n🧪 SANDBOX" if APP_ENV == "sandbox" else ""

        if APP_ENV == "sandbox":
            if not is_admin:
                await update.message.reply_text("❌ Доступ запрещён")
                return
            await update.message.reply_text(
                "🛠 Админ-панель системы" + env_marker,
                reply_markup=ReplyKeyboardRemove()
            )
            await self.cmd_management(update, context)
            return

        await update.message.reply_text(
            "👋 Добро пожаловать в News Aggregator Bot!" + env_marker + "\n\n"
            "Используйте /help для списка команд",
            reply_markup=self.REPLY_KEYBOARD
        )
    
    async def cmd_help(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /help"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        if get_app_env() == "sandbox":
            await update.message.reply_text(
                "🛠 Админ-режим\n\n"
                "Используйте кнопки админ-панели для управления системой.",
                reply_markup=self._build_sandbox_admin_keyboard(),
            )
            return
        help_text = (
            "📚 Доступные команды:\n\n"
            "/sync - Принудительно запустить сбор новостей\n"
            "/status - Показать статус бота и статистику\n"
            "/pause - Приостановить автоматический сбор\n"
            "/resume - Возобновить автоматический сбор\n"
            "/help - Показать эту справку\n\n"
            "⚙️ Нажмите кнопку 'Настройки' внизу для доступа к:\n"
            "  • Фильтр по категориям\n"
            "  • Управление источниками новостей\n\n"
            "Бот автоматически проверяет новости каждые 2 минуты"
        )
        await update.message.reply_text(help_text, reply_markup=self.REPLY_KEYBOARD)
    
    async def cmd_sync(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /sync - принудительный сбор новостей"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        user_id = update.message.from_user.id
        args = [arg.strip().lower() for arg in (context.args or [])]
        force = any(arg in ("force", "-f", "--force") for arg in args)
        if force and not self._is_admin(user_id):
            await update.message.reply_text("❌ Параметр force доступен только администраторам")
            return
        stop_state = get_global_collection_stop_state(app_env=get_app_env())
        if stop_state.enabled:
            ttl = stop_state.ttl_sec_remaining
            ttl_text = f" (TTL: {ttl}s)" if ttl is not None else ""
            await update.message.reply_text(f"⛔️ Сбор остановлен глобально{ttl_text}. /status покажет статус.")
            return

        await update.message.reply_text(
            "🔄 Начинаю сбор новостей..." if not force else "🔄 Начинаю сбор новостей (force)..."
        )
        
        try:
            count = await self.collect_and_publish(force=force)
            await update.message.reply_text(f"✅ Собрано и опубликовано {count} новостей")
        except Exception as e:
            logger.error(f"Error in sync: {e}")
            await update.message.reply_text(f"❌ Ошибка при сборе: {e}")
    
    async def cmd_debug_sources(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /debug_sources - показать все источники в БД"""
        if not self._is_admin(update.message.from_user.id):
            await update.message.reply_text("❌ Доступно только администраторам")
            return
        
        all_sources = self.db.get_all_sources()
        if not all_sources:
            await update.message.reply_text("📭 В БД нет новостей ни от одного источника")
            return
        
        text = "📋 Все источники в БД:\n\n"
        total = 0
        for source, count in all_sources.items():
            text += f"• {source}: {count}\n"
            total += count
        text += f"\n📊 Всего новостей: {total}"
        await update.message.reply_text(text)
    
    async def cmd_my_selection(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /my_selection - показать выбранные новости и экспортировать"""
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        user_id = update.message.from_user.id
        selected = self.db.get_user_selections(user_id, env="prod")
        
        if not selected:
            await update.message.reply_text("📭 У вас нет выбранных новостей.\n\nВыберите новости, нажав 📌 под новостью в канале.")
            return
        
        # Показать количество и кнопки действий
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("📄 Экспорт в DOC", callback_data="export_doc")],
            [InlineKeyboardButton("🗑 Очистить выбранное", callback_data="clear_selection")]
        ])
        
        await update.message.reply_text(
            f"📌 Выбрано новостей: {len(selected)}\n\n"
            f"Нажмите кнопку ниже для экспорта в документ.",
            reply_markup=keyboard
        )

    def _get_configured_source_maps(self) -> tuple[dict, dict, dict]:
        type_map: dict[str, str] = {}
        label_map: dict[str, str] = {}
        group_map: dict[str, str] = {}
        for entry in self.collector._configured_sources:
            if isinstance(entry, dict):
                fetch_url = entry.get('fetch_url', '')
                source_name = entry.get('source_name', '')
                src_type = entry.get('src_type', '')
            else:
                try:
                    fetch_url, source_name, _category, src_type = entry
                except ValueError:
                    continue

            if not source_name or source_name in type_map:
                continue

            source_type = src_type
            group = 'site'
            if '/telegram/channel/' in fetch_url:
                source_type = 'api'
                group = 'telegram'
            elif '/twitter/user/' in fetch_url:
                source_type = 'x/twitter'
            elif source_name in ('news.yahoo.com', 'rss.news.yahoo.com'):
                source_type = 'yahoo'
            elif src_type == 'rss':
                source_type = 'rss'
            else:
                source_type = 'html'

            label = f"@{source_name}" if group == 'telegram' else source_name
            type_map[source_name] = source_type
            label_map[source_name] = label
            group_map[source_name] = group
        return type_map, label_map, group_map

    def _build_source_status_sections(self, window_hours: int = 24) -> tuple[str, str]:
        type_map, label_map, group_map = self._get_configured_source_maps()
        sources = sorted(type_map.keys())
        if not sources:
            return "", ""

        counts = self.db.get_source_event_counts(sources, window_hours=window_hours)
        drop_codes = self.db.get_source_last_drop_codes(sources, window_hours=window_hours)
        health = self.db.get_source_health_snapshot(sources)
        error_rate_threshold = 0.5

        def _format_lines(group: str, title: str) -> str:
            lines = []
            for source in sorted(s for s in sources if group_map.get(s) == group):
                info = counts.get(source, {})
                success = info.get('success_count', 0)
                error = info.get('error_count', 0)
                drop_old = info.get('drop_old_count', 0)
                drop_date = info.get('drop_date_count', 0)
                total = success + error
                error_rate = (error / total) if total > 0 else 0.0

                is_green = success > 0 and error_rate < error_rate_threshold
                if is_green:
                    icon = "🟢"
                    status = f"{success} новости (24ч)"
                else:
                    if error > 0:
                        icon = "🔴"
                        status = health.get(source, {}).get('last_error_code') or "FETCH_ERROR"
                    elif drop_old > 0:
                        icon = "🟡"
                        status = "OLD_PUBLISHED_AT"
                    elif drop_date > 0:
                        icon = "🟡"
                        status = drop_codes.get(source) or "NO_PUBLISHED_DATE"
                    else:
                        icon = "🟡"
                        status = "FETCH_OK_NO_MATCH"

                source_type = type_map.get(source, 'rss')
                label = label_map.get(source, source)
                lines.append(f"{icon} {label} [{source_type}] — {status}")

            if not lines:
                return ""
            return f"\n{title}:\n" + "\n".join(lines) + "\n"

        channels_text = _format_lines('telegram', '📡 Каналы Telegram')
        sites_text = _format_lines('site', '🌐 Источники')
        return channels_text, sites_text

    def _build_status_text(self) -> str:
        stats = self.db.get_stats()
        ai_usage = self.db.get_ai_usage()
        channels_text, sites_text = self._build_source_status_sections(window_hours=24)

        input_tokens = int(ai_usage['total_tokens'] * 0.6)
        output_tokens = int(ai_usage['total_tokens'] * 0.4)
        input_cost = (input_tokens / 1_000_000.0) * 0.14
        output_cost = (output_tokens / 1_000_000.0) * 0.28
        estimated_cost = input_cost + output_cost

        daily_calls = 0
        daily_tokens = 0
        daily_cost = 0.0
        cache_hit_rate = 0.0
        budget_state = "OK"
        degraded_features = []

        if self.deepseek_client.budget:
            try:
                budget_state_data = self.deepseek_client.budget.get_state()
                budget_state = budget_state_data.get("budget_state", "OK")
                degraded_features = budget_state_data.get("degraded_features", [])
                daily = budget_state_data.get("usage", {})
                daily_calls = int(daily.get("calls", 0) or 0)
                daily_tokens = int((daily.get("tokens_in", 0) or 0) + (daily.get("tokens_out", 0) or 0))
                daily_cost = float(daily.get("cost_usd", 0.0) or 0.0)
                cache_hits = int(daily.get("cache_hits", 0) or 0)
                cache_hit_rate = (cache_hits / daily_calls * 100.0) if daily_calls > 0 else 0.0
            except Exception as e:
                logger.error(f"Error getting budget info: {e}")

        tick_state = self._get_ai_tick_state()
        gate_disabled = tick_state.get("disabled", [])
        combined_degraded = sorted(set(degraded_features + gate_disabled))
        degraded_text = ", ".join(combined_degraded) if combined_degraded else "-"

        stop_state = get_global_collection_stop_state(app_env=get_app_env())
        stop_label = "ON" if stop_state.enabled else "OFF"
        stop_ttl = stop_state.ttl_sec_remaining
        stop_ttl_text = f"{stop_ttl}s" if stop_ttl is not None else "-"

        status_text = (
            f"📊 Статус бота:\n\n"
            f"Статус: {'⏸️ PAUSED' if self.is_paused else '✅ RUNNING'}\n"
            f"Global stop: {stop_label} (TTL: {stop_ttl_text})\n"
            f"Всего опубликовано: {stats['total']}\n"
            f"За сегодня: {stats['today']}\n"
            f"Интервал проверки: {CHECK_INTERVAL_SECONDS} сек\n"
            f"───────────────────────────────\n"
            f"🧠 ИИ использование (всего):\n"
            f"Всего запросов: {ai_usage['total_requests']}\n"
            f"Всего токенов: {ai_usage['total_tokens']:,}\n"
            f"Расчетная стоимость: ${estimated_cost:.4f}\n\n"
            f"📝 Пересказы: {ai_usage['summarize_requests']} запр., {ai_usage['summarize_tokens']:,} токенов\n"
            f"🏷️ Категории: {ai_usage['category_requests']} запр., {ai_usage['category_tokens']:,} токенов\n"
            f"✨ Очистка текста: {ai_usage['text_clean_requests']} запр., {ai_usage['text_clean_tokens']:,} токенов\n"
            f"───────────────────────────────\n"
            f"📈 ИИ сегодня:\n"
            f"Вызовы: {daily_calls}\n"
            f"Токены: {daily_tokens:,}\n"
            f"Стоимость: ${daily_cost:.4f}\n"
            f"Cache hit rate: {cache_hit_rate:.1f}%\n"
            f"Вызовы в тике: {tick_state.get('calls', 0)}\n"
            f"Budget state: {budget_state}\n"
            f"Degraded: {degraded_text}\n"
            f"───────────────────────────────"
            f"{channels_text}"
            f"───────────────────────────────"
            f"{sites_text}"
        )
        return status_text
    
    async def cmd_status(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /status"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        status_text = self._build_status_text()
        await update.message.reply_text(status_text, disable_web_page_preview=True)
    
    async def cmd_pause(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /pause - приостановить новости для пользователя"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        if get_app_env() == "sandbox" and not self._is_admin(update.message.from_user.id):
            await update.message.reply_text("⛔ Access denied")
            return
        user_id = update.message.from_user.id
        self.db.set_pause_state(str(user_id), True, env="prod")
        logger.info(f"USER_PAUSE_SET user_id={user_id}")
        await update.message.reply_text("⏸️ Новости приостановлены для вас\n\nСбор продолжается, но вы не получаете уведомления.\nНажмите ▶️ для возобновления.")
    
    async def cmd_resume(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /resume - возобновить новости для пользователя"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        if get_app_env() == "sandbox" and not self._is_admin(update.message.from_user.id):
            await update.message.reply_text("⛔ Access denied")
            return
        user_id = update.message.from_user.id
        self.db.set_pause_state(str(user_id), False, env="prod")
        logger.info(f"USER_RESUME_SET user_id={user_id}")
        await update.message.reply_text("▶️ Новости возобновлены!\n\nТеперь вы снова получаете уведомления о новостях.")
        await self._deliver_pending_for_user(user_id)
    
    async def cmd_management(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """🛠 Management menu (sandbox admin only)"""
        try:
            from config.railway_config import APP_ENV
        except (ImportError, ValueError):
            from config.config import APP_ENV
        
        user_id = update.message.from_user.id
        
        # Check if sandbox and admin
        if APP_ENV != "sandbox":
            await update.message.reply_text("❌ Management available only in sandbox")
            return
        
        is_admin = self._is_admin(user_id)
        if not is_admin:
            await update.message.reply_text("❌ Доступно только администраторам")
            return

        # Show expanded management menu with all admin panels
        reply_markup = self._build_sandbox_admin_keyboard()
        await update.message.reply_text(
            "🛠 Управление системой",
            reply_markup=reply_markup,
            reply_to_message_id=update.message.message_id
        )
    
    async def cmd_sync_deepseek(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /sync_deepseek - показать текущую статистику и инструкцию"""
        ai_usage = self.db.get_ai_usage()
        
        # Calculate costs
        input_tokens = int(ai_usage['total_tokens'] * 0.6)
        output_tokens = int(ai_usage['total_tokens'] * 0.4)
        input_cost = (input_tokens / 1_000_000.0) * 0.14
        output_cost = (output_tokens / 1_000_000.0) * 0.28
        estimated_cost = input_cost + output_cost
        
        text = (
            f"📊 Текущая статистика в боте:\n\n"
            f"Запросов: {ai_usage['total_requests']}\n"
            f"Токенов: {ai_usage['total_tokens']:,}\n"
            f"Стоимость: ${estimated_cost:.4f}\n\n"
            f"🔄 Для синхронизации с реальными данными DeepSeek:\n\n"
            f"1️⃣ Откройте https://platform.deepseek.com/usage\n"
            f"2️⃣ Посмотрите данные:\n"
            f"   • API requests\n"
            f"   • Tokens\n" 
            f"   • Monthly expenses\n\n"
            f"3️⃣ Отправьте команду:\n"
            f"/update_stats <requests> <tokens> <cost>\n\n"
            f"Пример:\n"
            f"/update_stats 1331 413515 0.04"
        )
        await update.message.reply_text(text)
    
    async def cmd_update_stats(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /update_stats - синхронизировать с реальными данными DeepSeek"""
        try:
            # Если нет аргументов - показать текущие данные и инструкцию
            if not context.args or len(context.args) < 3:
                current = self.db.get_ai_usage()
                await update.message.reply_text(
                    f"📊 Текущие данные в боте:\n\n"
                    f"Запросов: {current['total_requests']}\n"
                    f"Токенов: {current['total_tokens']:,}\n"
                    f"Стоимость: ${current['total_cost_usd']:.4f}\n\n"
                    f"🔄 Для синхронизации используйте:\n"
                    f"/update_stats <requests> <tokens> <cost>\n\n"
                    f"Пример:\n"
                    f"/update_stats 1661 515627 0.06\n\n"
                    f"⚠️ Данные берите из DeepSeek:\n"
                    f"https://platform.deepseek.com/usage"
                )
                return
            
            requests = int(context.args[0])
            tokens = int(context.args[1])
            cost = float(context.args[2])
            
            # Get current stats
            current = self.db.get_ai_usage()
            
            # Use new sync method to set absolute values
            success = self.db.sync_ai_usage_with_deepseek(requests, tokens, cost)
            
            if success:
                await update.message.reply_text(
                    f"✅ Синхронизировано с DeepSeek!\n\n"
                    f"Было:\n"
                    f"📊 {current['total_requests']} → {requests} запросов\n"
                    f"🔢 {current['total_tokens']:,} → {tokens:,} токенов\n"
                    f"💰 ${current['total_cost_usd']:.4f} → ${cost:.4f}\n\n"
                    f"✨ Дальше учет идет автоматически!\n"
                    f"📈 Эти данные сохраняются и НЕ сбрасываются"
                )
            else:
                await update.message.reply_text("❌ Ошибка при синхронизации")
                
        except ValueError:
            await update.message.reply_text(
                "❌ Ошибка формата! Используйте числа.\n\n"
                "Пример: /update_stats 1661 515627 0.06"
            )
        except Exception as e:
            logger.error(f"Error updating stats: {e}")
            await update.message.reply_text(f"❌ Ошибка: {e}")
    
    async def cmd_filter(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /filter"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        # Создаем inline кнопки для выбора категорий
        ai_status = "✅" if self.ai_verification_enabled else "❌"
        keyboard = [
            [
                InlineKeyboardButton("#Мир", callback_data="filter_world"),
                InlineKeyboardButton("#Россия", callback_data="filter_russia"),
            ],
            [
                InlineKeyboardButton("#Москва", callback_data="filter_moscow"),
                InlineKeyboardButton("#Подмосковье", callback_data="filter_moscow_region"),
                InlineKeyboardButton("Все новости", callback_data="filter_all"),
            ],
            [
                InlineKeyboardButton(f"AI {ai_status}", callback_data="toggle_ai"),
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        ai_status_text = "включена" if self.ai_verification_enabled else "отключена"
        await update.message.reply_text(
            "Выберите категорию для фильтрации новостей в канале:\n\n"
            "#Мир - Новости со всего мира\n"
            "#Россия - Новости России\n"
            "#Москва - Новости Москвы\n"
            "#Подмосковье - Новости Московской области\n"
            "Все новости - Показывать все\n\n"
            f"🤖 AI верификация: {ai_status_text}",
            reply_markup=reply_markup
        )
    
    async def handle_emoji_buttons(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик эмодзи-кнопок"""
        if not await self._sandbox_admin_guard(update=update):
            return
        text = update.message.text
        user_id = update.message.from_user.id

        # Check if waiting for invite recipient
        if not hasattr(self, '_pending_invites'):
            self._pending_invites = {}
        
        # Old invite text input handler removed - now using share buttons

        # Custom export period input (hours)
        if context.user_data.get("awaiting_export_hours"):
            raw = (text or "").strip()
            try:
                hours = int(raw)
                if hours < 1 or hours > 24:
                    raise ValueError("hours out of range")
            except Exception:
                await update.message.reply_text(
                    "❌ Укажите число часов от 1 до 24.\n"
                    "Пример: 4"
                )
                return

            context.user_data["awaiting_export_hours"] = False
            await self._export_news_period(update.effective_user.id, context, hours=hours)
            return

        if context.user_data.get("awaiting_invite_label"):
            raw_label = (text or "").strip()
            context.user_data["awaiting_invite_label"] = False
            label = None if raw_label.lower() in ("пропустить", "skip", "-") else raw_label
            await self._finalize_invite_creation(
                admin_id=str(user_id),
                label=label,
                context=context,
                update=update,
            )
            return
        
        if get_app_env() == "sandbox" and text in {'🔄', '✉️', '⏸️', '▶️'}:
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return

        if text == '🔄':
            await self.cmd_sync(update, context)
        elif text == '✉️':
            # Отправить в личку (Мои новости)
            await self.cmd_my_selection(update, context)
        elif text == '⏸️':
            await self.cmd_pause(update, context)
        elif text == '▶️':
            await self.cmd_resume(update, context)
        elif text == '⚙️ Настройки':
            await self.cmd_settings(update, context)
        elif text == '🛠 Управление':
            await self.cmd_management(update, context)
    
    async def cmd_settings(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """⚙️ Меню настроек"""
        from core.services.global_stop import get_global_stop
        if get_app_env() == "sandbox":
            await update.message.reply_text("⛔ Недоступно в админ-режиме")
            return
        if get_app_env() == "prod" and get_global_stop():
            await update.message.reply_text("🔴 Система временно остановлена администратором.")
            return
        user_id = update.message.from_user.id
        is_admin = self._is_admin(user_id)
        app_env = get_app_env()

        keyboard = []
        keyboard.append([InlineKeyboardButton("🧰 Фильтр", callback_data="settings:filter")])
        
        # PROD mode: only user-friendly buttons
        if app_env == "prod":
            translate_enabled, target_lang = self.db.get_user_translation(str(user_id), env="prod")
            translate_status = "Вкл" if translate_enabled else "Выкл"
            delivery_mode = self.db.get_user_delivery_mode(str(user_id), env="prod")
            delivery_label = {"realtime": "В реальном времени", "hourly": "Каждый час", "morning": "Утром"}.get(delivery_mode, "В реальном времени")
            keyboard.append([InlineKeyboardButton("📰 Источники", callback_data="settings:sources:0")])
            keyboard.append([InlineKeyboardButton(f"🌐 Перевод ({target_lang.upper()}): {translate_status}", callback_data="settings:translate_toggle")])
            keyboard.append([InlineKeyboardButton(f"⏰ Доставка: {delivery_label}", callback_data="settings:delivery_mode")])
            keyboard.append([InlineKeyboardButton("📥 Экспорт новостей", callback_data="export_menu")])
            keyboard.append([InlineKeyboardButton("📊 Статус бота", callback_data="show_status")])
        else:
            # SANDBOX mode: include admin features
            keyboard.append([InlineKeyboardButton("🤖 AI переключатели", callback_data="ai:management")])
            keyboard.append([InlineKeyboardButton("📊 Статус бота", callback_data="show_status")])
            
            # Global collection control buttons for sandbox admins
            if is_admin:
                is_stopped, _ttl = get_global_collection_stop_status(app_env=app_env)
                if is_stopped:
                    keyboard.append([InlineKeyboardButton("▶️ Возобновить сбор", callback_data="collection:restore")])
                else:
                    keyboard.append([InlineKeyboardButton("⏸ Остановить сбор", callback_data="collection:stop")])

        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text("⚙️ Настройки", reply_markup=reply_markup)
    
    async def cmd_filter(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Команда /filter - выбор категорий для фильтрации"""
        # Создаем inline кнопки для выбора категорий
        ai_status = "✅" if self.ai_verification_enabled else "❌"
        app_env = get_app_env()
        
        selection_count = 0
        if app_env == "prod":
            user_id = update.message.from_user.id
            selection_count = len(self.db.get_user_selections(user_id, env="prod"))
        
        keyboard = [
            [
                InlineKeyboardButton("#Мир", callback_data="filter_world"),
                InlineKeyboardButton("#Россия", callback_data="filter_russia"),
            ],
            [
                InlineKeyboardButton("#Москва", callback_data="filter_moscow"),
                InlineKeyboardButton("#Подмосковье", callback_data="filter_moscow_region"),
                InlineKeyboardButton("Все новости", callback_data="filter_all"),
            ],
            [
                InlineKeyboardButton(f"AI {ai_status}", callback_data="toggle_ai"),
            ],
            [InlineKeyboardButton("📊 Статус бота", callback_data="show_status")],
        ]
        if app_env == "prod":
            keyboard.append([InlineKeyboardButton("📥 Unload", callback_data="export_menu")])
            keyboard.append([InlineKeyboardButton(f"📄 Мои новости ({selection_count})", callback_data="show_my_selection")])
        reply_markup = InlineKeyboardMarkup(keyboard)
        ai_status_text = "включена" if self.ai_verification_enabled else "отключена"
        await update.message.reply_text(
            "Выберите категорию для фильтрации новостей в канале:\n\n"
            "#Мир - Новости со всего мира\n"
            "#Россия - Новости России\n"
            "#Москва - Новости Москвы\n"
            "#Подмосковье - Новости Московской области\n"
            "Все новости - Показывать все\n\n"
            f"🤖 AI верификация: {ai_status_text}",
            reply_markup=reply_markup
        )
    
    async def button_callback(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        """Обработчик нажатия на кнопку"""
        query = update.callback_query

        if not await self._sandbox_admin_guard(query=query):
            return
        app_env = get_app_env()
        
        # ==================== PROD MODE RESTRICTIONS ====================
        # Block admin-only callbacks in prod environment
        if app_env == "prod":
            data = query.data or ""
            if data == "collection:stop" or data == "collection:restore":
                await query.answer(
                    "⛔ Остановка сбора доступна только в sandbox режиме",
                    show_alert=True
                )
                return
            if data == "mgmt:ai" or data == "ai:management" or data.startswith("mgmt:ai:"):
                await query.answer(
                    "⛔ AI-управление доступно только в sandbox режиме",
                    show_alert=True
                )
                return
        
        if app_env == "sandbox":
            data = query.data or ""
            if (
                data.startswith("settings:sources:")
                or data.startswith("settings:src_toggle:")
                or data.startswith("settings:src_page:")
                or data == "settings:translate_toggle"
                or data == "settings:delivery_mode"
                or data.startswith("settings:delivery:")
                or data == "export_menu"
                or data.startswith("export_period:")
                or data == "export_doc"
                or data == "clear_selection"
                or data == "show_my_selection"
                or data.startswith("select:")
            ):
                await query.answer("⛔ Access denied", show_alert=True)
                return
        
        # ==================== COLLECTION CONTROL CALLBACKS ====================
        if query.data == "collection:stop":
            # Stop global collection
            await query.answer()
            user_id = query.from_user.id
            if not self._is_admin(user_id):
                await query.edit_message_text("❌ Только администраторы могут остановить сбор")
                return

            set_global_collection_stop(True, ttl_sec=3600, by=str(user_id))
            
            # Уведомляем asyncio.Event об остановке (локальный эффект)
            from core.services.global_stop import set_global_stop
            set_global_stop(True)
            
            await query.edit_message_text(
                "⏸ Сбор новостей остановлен глобально\n\n"
                "Все боты перестали собирать новости.\n"
                "Используйте кнопку Восстановить для запуска.",
                reply_markup=InlineKeyboardMarkup([[
                    InlineKeyboardButton("▶️ Возобновить сбор", callback_data="collection:restore")
                ]])
            )
            return
        
        if query.data == "collection:restore":
            # Restore global collection
            await query.answer()
            user_id = query.from_user.id
            if not self._is_admin(user_id):
                await query.edit_message_text("❌ Только администраторы могут восстановить сбор")
                return

            set_global_collection_stop(False, by=str(user_id))
            from core.services.global_stop import set_global_stop
            set_global_stop(False)
            await query.edit_message_text(
                "▶️ Сбор новостей восстановлен!\n\n"
                "Боты снова собирают новости в фоне."
            )
            return
        
        # ==================== SETTINGS CALLBACKS ====================
        if query.data == "settings:filter":
            # Показать меню фильтра
            await query.answer()
            ai_status = "✅" if self.ai_verification_enabled else "❌"
            keyboard = [
                [
                    InlineKeyboardButton("#Мир", callback_data="filter_world"),
                    InlineKeyboardButton("#Россия", callback_data="filter_russia"),
                ],
                [
                    InlineKeyboardButton("#Москва", callback_data="filter_moscow"),
                    InlineKeyboardButton("#Подмосковье", callback_data="filter_moscow_region"),
                    InlineKeyboardButton("Все новости", callback_data="filter_all"),
                ],
                [
                    InlineKeyboardButton(f"AI {ai_status}", callback_data="toggle_ai"),
                ],
                [
                    InlineKeyboardButton("⬅️ Назад", callback_data="settings:back"),
                ]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            ai_status_text = "включена" if self.ai_verification_enabled else "отключена"
            await query.edit_message_text(
                text=(
                    "🧰 Фильтр\n\n"
                    "#Мир - Новости со всего мира\n"
                    "#Россия - Новости России\n"
                    "#Москва - Новости Москвы\n"
                    "#Подмосковье - Новости Московской области\n"
                    "Все новости - Показывать все\n\n"
                    f"🤖 AI верификация: {ai_status_text}"
                ),
                reply_markup=reply_markup
            )
            return
        
        if query.data.startswith("settings:sources:"):
            # Показать список источников
            await query.answer()
            page = int(query.data.split(":")[-1])
            await self._show_sources_menu(query, page)
            return
        
        if query.data.startswith("settings:src_toggle:"):
            # Переключить источник
            parts = query.data.split(":")
            source_id = int(parts[2])
            page = int(parts[3]) if len(parts) > 3 else 0
            
            user_id = query.from_user.id
            new_state = self.db.toggle_user_source(user_id, source_id, env="prod")
            
            await query.answer(f"{'✅ Включено' if new_state else '❌ Отключено'}", show_alert=False)
            await self._show_sources_menu(query, page)
            return
        
        if query.data.startswith("settings:src_page:"):
            # Пагинация источников
            page = int(query.data.split(":")[-1])
            await query.answer()
            await self._show_sources_menu(query, page)
            return
        
        if query.data == "settings:back":
            # Вернуться к меню настроек
            await query.answer()
            user_id = str(query.from_user.id)
            translate_enabled, target_lang = self.db.get_user_translation(user_id, env="prod")
            translate_status = "Вкл" if translate_enabled else "Выкл"
            delivery_mode = self.db.get_user_delivery_mode(user_id, env="prod")
            delivery_label = {"realtime": "В реальном времени", "hourly": "Каждый час", "morning": "Утром"}.get(delivery_mode, "В реальном времени")
            keyboard = [
                [InlineKeyboardButton("🧰 Фильтр", callback_data="settings:filter")],
                [InlineKeyboardButton("📰 Источники", callback_data="settings:sources:0")],
                [InlineKeyboardButton(f"🌐 Перевод ({target_lang.upper()}): {translate_status}", callback_data="settings:translate_toggle")],
                [InlineKeyboardButton(f"⏰ Доставка: {delivery_label}", callback_data="settings:delivery_mode")],
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                text="⚙️ Настройки",
                reply_markup=reply_markup
            )
            return

        if query.data == "settings:translate_toggle":
            await query.answer()
            user_id = str(query.from_user.id)
            enabled, target_lang = self.db.get_user_translation(user_id, env="prod")
            new_enabled = not enabled
            self.db.set_user_translation(user_id, new_enabled, target_lang, env="prod")

            status_text = "Включен" if new_enabled else "Выключен"
            await query.edit_message_text(
                text=f"🌐 Перевод ({target_lang.upper()}) {status_text}\n\nПеревод применяется к англоязычным новостям.",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")]
                ])
            )
            return
        
        if query.data == "settings:delivery_mode":
            await query.answer()
            user_id = str(query.from_user.id)
            current_mode = self.db.get_user_delivery_mode(user_id, env="prod")
            
            keyboard = [
                [InlineKeyboardButton(
                    f"{'✅' if current_mode == 'realtime' else '⚪'} В реальном времени",
                    callback_data="settings:delivery:realtime"
                )],
                [InlineKeyboardButton(
                    f"{'✅' if current_mode == 'hourly' else '⚪'} Каждый час",
                    callback_data="settings:delivery:hourly"
                )],
                [InlineKeyboardButton(
                    f"{'✅' if current_mode == 'morning' else '⚪'} Утром (7:00)",
                    callback_data="settings:delivery:morning"
                )],
                [InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")]
            ]
            
            await query.edit_message_text(
                text="⏰ Режим доставки новостей:\n\n"
                     "• В реальном времени — новости приходят сразу\n"
                     "• Каждый час — новости собираются и отправляются раз в час\n"
                     "• Утром — новости приходят один раз в 7:00",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
            return
        
        if query.data.startswith("settings:delivery:"):
            await query.answer()
            user_id = str(query.from_user.id)
            mode = query.data.split(":")[-1]  # realtime, hourly, morning
            
            success = self.db.set_user_delivery_mode(user_id, mode, env="prod")
            if success:
                mode_names = {
                    "realtime": "В реальном времени",
                    "hourly": "Каждый час",
                    "morning": "Утром (7:00)"
                }
                await query.edit_message_text(
                    text=f"✅ Режим доставки изменен на: {mode_names.get(mode, mode)}",
                    reply_markup=InlineKeyboardMarkup([
                        [InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")]
                    ])
                )
            else:
                await query.edit_message_text(
                    text="❌ Ошибка при изменении режима доставки",
                    reply_markup=InlineKeyboardMarkup([
                        [InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")]
                    ])
                )
            return
        
        # ==================== AI MANAGEMENT CALLBACKS (ALL ADMINS) ====================
        if query.data == "ai:management":
            # Show AI levels management screen (works on prod too)
            await query.answer()
            await self._show_ai_management(query)
            return
        
        if query.data.startswith("ai:inc:"):
            # Increment AI level
            module = query.data.split(":")[-1]
            await self._handle_ai_level_change(query, module, action="inc")
            return
        
        elif query.data.startswith("filter_"):
            # Фильтрация по категориям
            filter_type = query.data.replace("filter_", "")
            new_filter = filter_type if filter_type != 'all' else None
            app_env = get_app_env()
            if app_env == "sandbox":
                self._set_global_category_filter(new_filter)
            else:
                self._set_user_category_filter(query.from_user.id, new_filter)

            filter_names = {
                'world': '#Мир',
                'russia': '#Россия',
                'moscow': '#Москва',
                'moscow_region': '#Подмосковье',
                'all': 'Все новости'
            }

            await query.answer(
                f"✅ Фильтр установлен: {filter_names.get(filter_type, 'Неизвестно')}",
                show_alert=False,
            )
            await query.edit_message_text(
                text=f"✅ Установлена фильтрация: {filter_names.get(filter_type, 'Неизвестно')}\n\n"
                     f"Режим: {'глобальный (sandbox)' if app_env == 'sandbox' else 'персональный (prod)'}"
            )
            return
        # ==================== MANAGEMENT CALLBACKS (SANDBOX ADMIN ONLY) ====================
        if query.data.startswith("mgmt:"):
            if get_app_env() != "sandbox" and not query.data.startswith("mgmt:send_invite:"):
                await query.answer("❌ Управление доступно только в песочнице", show_alert=True)
                return
        if query.data.startswith("mgmt:send_invite:"):
            # Show share options for invite (works in sandbox only)
            await query.answer()
            try:
                from config.railway_config import APP_ENV
            except (ImportError, ValueError):
                from config.config import APP_ENV
            
            if APP_ENV != "sandbox":
                await query.edit_message_text("❌ Отправка инвайтов доступна только в песочнице")
                return
            
            # Extract invite code from callback data
            invite_code = query.data.split(":", 2)[2]
            logger.info(f"Preparing to share invite {invite_code}")
            invite_label = self.db.get_invite_label(invite_code)
            
            # Get PROD bot username (инвайт должен вести на прод бота)
            try:
                from config.railway_config import BOT_PROD_USERNAME
            except (ImportError, ValueError):
                try:
                    from config.config import BOT_PROD_USERNAME
                except ImportError:
                    BOT_PROD_USERNAME = "Tops_News_bot"  # Default prod bot
            
            if not BOT_PROD_USERNAME:
                BOT_PROD_USERNAME = "Tops_News_bot"
            
            # Формируем правильную ссылку на ПРОД бота
            invite_link = f"https://t.me/{BOT_PROD_USERNAME}?start={invite_code}"
            
            # Красивое сообщение с эмодзи (без ссылки на бота)
            from urllib.parse import quote
            if invite_label:
                from html import escape
                label_line = f"👤 Для: {escape(invite_label)}\n"
            else:
                label_line = ""
            share_text = quote(
                "🎁 Приглашение в News Aggregator Bot!\n\n"
                f"{label_line}"
                "✨ Используйте этот инвайт-код для регистрации:\n"
                f"👉 {invite_code}\n\n"
                f"🚀 Перейти: {invite_link}"
            )
            
            share_url = f"https://t.me/share/url?url={invite_link}&text={share_text}"
            
            keyboard = [
                [InlineKeyboardButton("📤 Поделиться инвайтом", url=share_url)],
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")],
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            
            await query.edit_message_text(
                text=(
                    f"📤 Отправка инвайта\n\n"
                    f"Нажмите кнопку 'Поделиться' и выберите контакт из Telegram\n\n"
                    f"{label_line}"
                    f"📌 Код инвайта: <code>{invite_code}</code>\n"
                    f"🔗 Ссылка: <code>{invite_link}</code>"
                ),
                reply_markup=reply_markup,
                parse_mode='HTML'
            )
            
            return
        
        if query.data == "mgmt:users":
            # Show users and invites management screen
            await query.answer()
            await self._show_users_management(query)
            return
        
        if query.data.startswith("mgmt:ai:dec:"):
            # Decrement AI level
            module = query.data.split(":")[-1]
            await self._handle_ai_level_change(query, module, action="dec")
            return
        
        if query.data.startswith("mgmt:ai:set:"):
            # Set AI level directly
            parts = query.data.split(":")
            module = parts[2]
            level = int(parts[3])
            await self._handle_ai_level_change(query, module, action="set", level=level)
            return
        
        if query.data == "mgmt:back":
            # Back to management main menu
            await query.answer()
            keyboard = [
                [InlineKeyboardButton("🤖 AI переключатели", callback_data="ai:management")],
                [InlineKeyboardButton("👥 Пользователи и инвайты", callback_data="mgmt:users")],
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                text="🛠 Управление системой",
                reply_markup=reply_markup
            )
        
            return
        
        if query.data == "mgmt:new_invite":
            context.user_data["awaiting_invite_label"] = True
            keyboard = [
                [InlineKeyboardButton("Пропустить", callback_data="mgmt:invite_label:skip")],
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")],
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(
                text=(
                    "✍️ Введите имя или комментарий для инвайта.\n"
                    "Например: Иван (редакция), Петр С.\n\n"
                    "Можно нажать 'Пропустить', если имя не нужно."
                ),
                reply_markup=reply_markup
            )
            return

        if query.data == "mgmt:invite_label:skip":
            await query.answer()
            context.user_data["awaiting_invite_label"] = False
            await self._finalize_invite_creation(
                admin_id=str(query.from_user.id),
                label=None,
                context=context,
                query=query,
            )
            return
        
        if query.data == "mgmt:users_list":
            # Show detailed list of users with block/unblock buttons
            approved_users = self.access_db.get_approved_users()
            
            if not approved_users:
                await query.edit_message_text(
                    text="✅ Список одобренных пользователей пуст",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")]])
                )
                return
            
            # Показываем по одному пользователю с кнопками действия
            # (Telegram имеет ограничение на размер сообщения)
            user_id, username, first_name, approved_at, invited_by, invite_label = approved_users[0]
            name = first_name or username or user_id
            
            text = f"👤 Управление пользователями\n\n"
            text += f"Пользователь: <b>{name}</b>\n"
            text += f"ID: <code>{user_id}</code>\n"
            text += f"Username: {f'@{username}' if username else 'нет'}\n"
            text += f"Одобрен: {approved_at}\n\n"
            if invite_label:
                from html import escape
                text += f"Инвайт: {escape(invite_label)}\n"
            if invited_by:
                text += f"Кем приглашен: {invited_by}\n"
            text += f"Всего одобренных: {len(approved_users)}\n"
            
            # Кнопки для управления
            keyboard = [
                [
                    InlineKeyboardButton("🔒 Заблокировать", callback_data=f"mgmt:block_user:{user_id}"),
                    InlineKeyboardButton("➡️ Далее", callback_data=f"mgmt:users_list_page:1")
                ],
                [InlineKeyboardButton("⬅️ Назад к управлению", callback_data="mgmt:users")]
            ]
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=text, reply_markup=reply_markup, parse_mode='HTML')
            return
        
        # Pagination for users list
        if query.data.startswith("mgmt:users_list_page:"):
            page = int(query.data.split(":")[2])
            approved_users = self.access_db.get_approved_users()
            
            if page >= len(approved_users):
                page = len(approved_users) - 1
            
            if page < 0 or not approved_users:
                await query.answer("Нет пользователей", show_alert=True)
                return
            
            user_id, username, first_name, approved_at, invited_by, invite_label = approved_users[page]
            name = first_name or username or user_id
            
            text = f"👤 Управление пользователями\n\n"
            text += f"Пользователь: <b>{name}</b>\n"
            text += f"ID: <code>{user_id}</code>\n"
            text += f"Username: {f'@{username}' if username else 'нет'}\n"
            text += f"Одобрен: {approved_at}\n\n"
            if invite_label:
                from html import escape
                text += f"Инвайт: {escape(invite_label)}\n"
            if invited_by:
                text += f"Кем приглашен: {invited_by}\n"
            text += f"Пользователь {page + 1} из {len(approved_users)}\n"
            
            # Navigation and action buttons
            keyboard = []
            keyboard.append([
                InlineKeyboardButton("🔒 Заблокировать", callback_data=f"mgmt:block_user:{user_id}")
            ])
            
            nav_buttons = []
            if page > 0:
                nav_buttons.append(InlineKeyboardButton("◀️ Назад", callback_data=f"mgmt:users_list_page:{page - 1}"))
            if page < len(approved_users) - 1:
                nav_buttons.append(InlineKeyboardButton("Далее ▶️", callback_data=f"mgmt:users_list_page:{page + 1}"))
            
            if nav_buttons:
                keyboard.append(nav_buttons)
            
            keyboard.append([InlineKeyboardButton("⬅️ Назад к управлению", callback_data="mgmt:users")])
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=text, reply_markup=reply_markup, parse_mode='HTML')
            return
        
        # Block user
        if query.data.startswith("mgmt:block_user:"):
            user_id = query.data.split(":")[2]
            if self.access_db.block_user(user_id):
                await query.answer(f"✅ Пользователь {user_id} заблокирован", show_alert=True)
                await query.edit_message_text(
                    text="✅ Пользователь успешно заблокирован",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")]])
                )
            else:
                await query.answer("❌ Ошибка при блокировке", show_alert=True)
            return
            
            await query.edit_message_text(text=text, reply_markup=reply_markup)
            return
        
        # Admin panel: System Status
        if query.data == "mgmt:status":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_status(query)
            return
        
        # Admin panel: AI Management
        if query.data == "mgmt:ai":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_ai_panel(query)
            return
        
        # Admin panel: Sources
        if query.data == "mgmt:sources":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_sources_panel(query)
            return
        
        # Admin panel: Statistics
        if query.data == "mgmt:stats":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_stats_panel(query)
            return
        
        # Admin panel: Settings
        if query.data == "mgmt:settings":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_settings_panel(query)
            return

        # Admin panel: Diagnostics
        if query.data == "mgmt:diag":
            await query.answer()
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await self._show_admin_diagnostics_panel(query)
            return
        
        # Back to admin menu
        if query.data == "mgmt:main":
            await query.answer()
            await self.cmd_management_inline(query)
            return
        
        # Toggle global stop
        if query.data == "mgmt:toggle_global_stop":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            from core.services.global_stop import toggle_global_stop
            new_state = toggle_global_stop()
            if new_state:
                logger.warning(f"[ADMIN] SYSTEM FULL STOP by {query.from_user.id}")
                await query.answer("🔴 Система полностью остановлена", show_alert=True)
            else:
                logger.warning(f"[ADMIN] SYSTEM FULL RESUME by {query.from_user.id}")
                await query.answer("🟢 Система возобновлена", show_alert=True)
            await query.edit_message_reply_markup(
                reply_markup=self._build_sandbox_admin_keyboard()
            )
            return
        
        # Resume work - запуск сбора новостей в песочнице
        if query.data == "mgmt:resume_work":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            
            # Получаем текущее окружение и настройки
            try:
                from config.railway_config import APP_ENV, CHECK_INTERVAL_SECONDS
            except (ImportError, ValueError):
                from config.config import APP_ENV, CHECK_INTERVAL_SECONDS
            
            # Проверяем текущее состояние системы ДО изменения
            from core.services.global_stop import get_global_stop, set_global_stop
            was_stopped = get_global_stop()
            
            # Снимаем глобальную остановку (если была установлена)
            set_global_stop(enabled=False, reason="Resume via admin button", by=f"admin_{query.from_user.id}")
            logger.warning(f"[ADMIN] WORK RESUMED by {query.from_user.id}, was_stopped={was_stopped}")
            
            # Формируем динамическое сообщение в зависимости от состояния
            interval_minutes = CHECK_INTERVAL_SECONDS // 60
            interval_seconds = CHECK_INTERVAL_SECONDS % 60
            
            if interval_seconds > 0:
                time_text = f"{interval_minutes} мин {interval_seconds} сек"
            else:
                time_text = f"{interval_minutes} мин"
            
            if was_stopped:
                # Система была остановлена администратором
                message = (
                    "🟢 Работа возобновлена!\n\n"
                    "✅ Остановка снята\n"
                    "📰 Сбор новостей начнется немедленно\n"
                    f"⏱ Периодичность: каждые {time_text}\n"
                    "🤖 AI модули активны"
                )
                # Запускаем сбор новостей в фоне только в sandbox
                if APP_ENV == "sandbox":
                    asyncio.create_task(self._trigger_news_collection())
            else:
                # Система уже работала
                message = (
                    "✅ Система работает!\n\n"
                    "🔄 Сбор новостей активен\n"
                    f"⏱ Периодичность: каждые {time_text}\n"
                    f"📰 Новости поступают каждые {time_text}\n"
                    "🤖 AI модули активны"
                )
            
            # Показываем модальное уведомление
            await query.answer(message, show_alert=True)
            
            # Обновляем клавиатуру
            await query.edit_message_reply_markup(
                reply_markup=self._build_sandbox_admin_keyboard()
            )
            return
        
        # AI module selection
        if query.data.startswith("mgmt:ai:module:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            module = query.data.split(":")[-1]
            await self._show_ai_module_control(query, module)
            return
        
        # AI level control buttons
        if query.data.startswith("mgmt:ai:level:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            parts = query.data.split(":")
            module = parts[2]
            action = parts[3]
            level = int(parts[4]) if len(parts) > 4 else 0
            
            if action == "inc":
                new_level = min(level + 1, 5)
            elif action == "dec":
                new_level = max(level - 1, 0)
            else:
                new_level = level
            
            # Save new level using AILevelManager
            from core.services.access_control import AILevelManager
            ai_manager = AILevelManager(self.db)
            ai_manager.set_level('global', module, new_level)
            
            await query.answer(f"✅ {module}: уровень {new_level}")
            await self._show_ai_module_control(query, module)
            return
        
        # Sources management
        if query.data == "mgmt:sources:toggle_all":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            # Toggle all sources
            logger.info(f"Sources toggle_all by admin_id={query.from_user.id}")
            await self._show_admin_sources_panel(query)
            return

        if query.data == "mgmt:sources:toggle_telegram":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            enabled = self._get_rsshub_telegram_enabled()
            new_value = "0" if enabled else "1"
            try:
                self.db.set_system_setting("rsshub_telegram_enabled", new_value)
            except Exception:
                pass
            await query.answer(
                "✅ Telegram RSSHub включен" if new_value == "1" else "⛔ Telegram RSSHub отключен",
                show_alert=True,
            )
            await self._show_admin_sources_panel(query)
            return
        
        if query.data == "mgmt:sources:rescan":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer("🔄 Переоценка источников запущена...", show_alert=False)
            logger.info(f"Sources rescan requested by admin_id={query.from_user.id}")
            await self._show_admin_sources_panel(query)
            return
        
        # Sources pagination
        if query.data.startswith("mgmt:sources:page:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            page = int(query.data.split(":")[-1])
            await self._show_admin_sources_panel(query, page=page)
            return
        
        # Source detail view
        if query.data.startswith("mgmt:source:detail:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            source_code = query.data.split(":")[-1]
            await self._show_source_detail(query, source_code)
            return
        
        # Source toggle enable/disable
        if query.data.startswith("mgmt:source:toggle:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            source_code = query.data.split(":")[-1]
            success = self.db.toggle_source_enabled(source_code)
            if success:
                await query.answer("✅ Статус изменен", show_alert=False)
            else:
                await query.answer("❌ Ошибка изменения статуса", show_alert=True)
            await self._show_source_detail(query, source_code)
            return
        
        # Source tier change
        if query.data.startswith("mgmt:source:tier:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            parts = query.data.split(":")
            source_code = parts[3]
            new_tier = parts[4]
            success = self.db.set_source_tier(source_code, new_tier)
            if success:
                await query.answer(f"✅ Tier изменен на {new_tier}", show_alert=False)
            else:
                await query.answer("❌ Ошибка изменения tier", show_alert=True)
            await self._show_source_detail(query, source_code)
            return
        
        # Source restore from quarantine
        if query.data.startswith("mgmt:source:restore:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            source_code = query.data.split(":")[-1]
            success = self.db.restore_source(source_code)
            if success:
                await query.answer("✅ Источник восстановлен", show_alert=True)
            else:
                await query.answer("❌ Ошибка восстановления", show_alert=True)
            await self._show_source_detail(query, source_code)
            return
        
        # Stats refresh
        if query.data == "mgmt:stats:refresh":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            await self._show_admin_stats_panel(query)
            return
        
        # Settings management
        if query.data == "mgmt:settings:interval":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            text = (
                "⏱️ ИНТЕРВАЛ ПРОВЕРКИ\n\n"
                "Текущее значение: 300 секунд\n"
                "Минимум: 60 секунд\n"
                "Максимум: 3600 секунд\n\n"
                "Выберите новое значение:"
            )
            keyboard = [
                [InlineKeyboardButton("60s", callback_data="mgmt:settings:interval:60"),
                 InlineKeyboardButton("120s", callback_data="mgmt:settings:interval:120"),
                 InlineKeyboardButton("300s", callback_data="mgmt:settings:interval:300")],
                [InlineKeyboardButton("600s", callback_data="mgmt:settings:interval:600"),
                 InlineKeyboardButton("1200s", callback_data="mgmt:settings:interval:1200")],
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:settings")],
            ]
            await query.edit_message_text(text=text, reply_markup=InlineKeyboardMarkup(keyboard))
            return
        
        if query.data.startswith("mgmt:settings:interval:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            interval = int(query.data.split(":")[-1])
            logger.info(f"CHECK_INTERVAL changed to {interval}s by admin_id={query.from_user.id}")
            await query.answer(f"✅ Интервал установлен на {interval}с", show_alert=True)
            await self._show_admin_settings_panel(query)
            return
        
        if query.data == "mgmt:settings:parallel":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            text = (
                "🔄 ПАРАЛЛЕЛЬНЫЕ ЗАДАЧИ\n\n"
                "Текущее значение: 3\n"
                "Минимум: 1\n"
                "Максимум: 10\n\n"
                "Выберите новое значение:"
            )
            keyboard = [
                [InlineKeyboardButton("1", callback_data="mgmt:settings:parallel:1"),
                 InlineKeyboardButton("2", callback_data="mgmt:settings:parallel:2"),
                 InlineKeyboardButton("3", callback_data="mgmt:settings:parallel:3")],
                [InlineKeyboardButton("5", callback_data="mgmt:settings:parallel:5"),
                 InlineKeyboardButton("10", callback_data="mgmt:settings:parallel:10")],
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:settings")],
            ]
            await query.edit_message_text(text=text, reply_markup=InlineKeyboardMarkup(keyboard))
            return
        
        if query.data.startswith("mgmt:settings:parallel:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            parallel = int(query.data.split(":")[-1])
            logger.info(f"Parallel tasks changed to {parallel} by admin_id={query.from_user.id}")
            await query.answer(f"✅ Параллельные задачи: {parallel}", show_alert=True)
            await self._show_admin_settings_panel(query)
            return
        
        if query.data == "mgmt:settings:logging":
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            await query.answer()
            text = (
                "📝 УРОВЕНЬ ЛОГИРОВАНИЯ\n\n"
                "Текущий уровень: INFO\n"
                "Выберите новый уровень:"
            )
            keyboard = [
                [InlineKeyboardButton("DEBUG", callback_data="mgmt:settings:logging:DEBUG"),
                 InlineKeyboardButton("INFO", callback_data="mgmt:settings:logging:INFO")],
                [InlineKeyboardButton("WARNING", callback_data="mgmt:settings:logging:WARNING"),
                 InlineKeyboardButton("ERROR", callback_data="mgmt:settings:logging:ERROR")],
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:settings")],
            ]
            await query.edit_message_text(text=text, reply_markup=InlineKeyboardMarkup(keyboard))
            return
        
        if query.data.startswith("mgmt:settings:logging:"):
            if not self._is_admin(query.from_user.id):
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return
            level = query.data.split(":")[-1]
            logger.info(f"Log level changed to {level} by admin_id={query.from_user.id}")
            await query.answer(f"✅ Уровень логирования: {level}", show_alert=True)
            await self._show_admin_settings_panel(query)
            return

        if query.data == "noop":
            await query.answer()
            return
        # ==================== OTHER CALLBACKS ====================
        if query.data == "show_status":
            # Показать статус бота
            await query.answer()
            user_id = query.from_user.id

            from core.services.global_stop import get_global_stop
            if get_app_env() == "prod" and get_global_stop():
                await context.bot.send_message(
                    chat_id=user_id,
                    text="🔴 Система временно остановлена администратором.",
                )
                return

            status_text = self._build_status_text()
            
            await context.bot.send_message(
                chat_id=user_id,
                text=status_text,
                disable_web_page_preview=True
            )
            return
        
        if query.data == "show_my_selection":
            # Показать выбранные новости с кнопками экспорта
            user_id = query.from_user.id
            selected = self.db.get_user_selections(user_id, env="prod")
            
            if not selected:
                await query.answer("📭 У вас нет выбранных новостей", show_alert=True)
                return
            
            keyboard = InlineKeyboardMarkup([
                [InlineKeyboardButton("📄 Экспорт в DOC", callback_data="export_doc")],
                [InlineKeyboardButton("🗑 Очистить выбранное", callback_data="clear_selection")]
            ])
            
            await query.answer()
            await context.bot.send_message(
                chat_id=user_id,
                text=f"📌 Выбрано новостей: {len(selected)}\n\nНажмите кнопку ниже для экспорта в документ.",
                reply_markup=keyboard
            )
            return

        if query.data == "export_menu":
            await query.answer()
            user_id = query.from_user.id

            keyboard = InlineKeyboardMarkup([
                [
                    InlineKeyboardButton("⏱ 1 час", callback_data="export_period:1"),
                    InlineKeyboardButton("⏱ 2 часа", callback_data="export_period:2"),
                    InlineKeyboardButton("⏱ 3 часа", callback_data="export_period:3"),
                ],
                [
                    InlineKeyboardButton("⏱ 6 часов", callback_data="export_period:6"),
                    InlineKeyboardButton("⏱ 12 часов", callback_data="export_period:12"),
                    InlineKeyboardButton("⏱ 24 часа", callback_data="export_period:24"),
                ],
                [
                    InlineKeyboardButton("🧩 Custom", callback_data="export_period:custom"),
                ]
            ])

            await context.bot.send_message(
                chat_id=user_id,
                text=(
                    "📥 Unload: выберите период выгрузки (макс. 24 часа).\n"
                    "Можно выбрать фиксированный период или Custom для своего значения."
                ),
                reply_markup=keyboard
            )
            return

        if query.data.startswith("export_period:"):
            await query.answer()
            period = query.data.split(":", 1)[1]
            user_id = query.from_user.id

            if period == "custom":
                context.user_data["awaiting_export_hours"] = True
                await context.bot.send_message(
                    chat_id=user_id,
                    text="🧩 Введите период в часах (1–24). Пример: 4"
                )
                return

            try:
                hours = int(period)
            except ValueError:
                await context.bot.send_message(chat_id=user_id, text="❌ Некорректный период")
                return

            await self._export_news_period(user_id, context, hours=hours)
            return
        
        if query.data == "export_doc":
            # Экспорт выбранных новостей в DOC
            user_id = query.from_user.id
            await query.answer("📄 Генерирую документ...", show_alert=False)
            
            try:
                doc_file = await self._generate_doc_file(user_id)
                if doc_file:
                    count = len(self.db.get_user_selections(user_id, env="prod"))
                    await context.bot.send_document(
                        chat_id=user_id,
                        document=open(doc_file, 'rb'),
                        filename="selected_news.docx",
                        caption=f"📰 Ваши выбранные новости ({count} шт.)"
                    )
                    # Удалить временный файл
                    import os
                    os.remove(doc_file)
                    
                    # Очистить выбранные новости после отправки
                    self.db.clear_user_selections(user_id, env="prod")
                    await context.bot.send_message(
                        chat_id=user_id,
                        text="✅ Документ отправлен!\n\n📌 Выбранные новости очищены. Начните новую подборку!"
                    )
                else:
                    await context.bot.send_message(user_id, "❌ Ошибка при создании документа")
            except Exception as e:
                logger.error(f"Error generating doc: {e}")
                await context.bot.send_message(user_id, f"❌ Ошибка: {str(e)[:100]}")
            return
        
        elif query.data == "clear_selection":
            # Очистить выбранные новости
            user_id = query.from_user.id
            count = len(self.db.get_user_selections(user_id, env="prod"))
            self.db.clear_user_selections(user_id, env="prod")
            await query.answer(f"🗑 Очищено {count} новостей", show_alert=False)
            await query.edit_message_text("✅ Выбранные новости очищены")
            return
        
        if query.data == "toggle_ai":
            # Переключение AI верификации
            self.ai_verification_enabled = not self.ai_verification_enabled
            status = "включена" if self.ai_verification_enabled else "отключена"
            emoji = "✅" if self.ai_verification_enabled else "❌"
            
            await query.answer(f"{emoji} AI верификация {status}", show_alert=False)
            await query.edit_message_text(
                text=f"{emoji} AI верификация категорий {status}\n\n"
                     f"DeepSeek {'теперь будет проверять' if self.ai_verification_enabled else 'больше не будет проверять'} "
                     "правильность определения категорий новостей."
            )
            return
        
        elif query.data.startswith("filter_"):
            # Фильтрация по категориям
            filter_type = query.data.replace("filter_", "")
            new_filter = filter_type if filter_type != 'all' else None
            app_env = get_app_env()
            if app_env == "sandbox":
                self._set_global_category_filter(new_filter)
            else:
                self._set_user_category_filter(query.from_user.id, new_filter)
            
            filter_names = {
                'world': '#Мир',
                'russia': '#Россия',
                'moscow': '#Москва',
                'moscow_region': '#Подмосковье',
                'all': 'Все новости'
            }
            
            await query.answer(f"✅ Фильтр установлен: {filter_names.get(filter_type, 'Неизвестно')}", show_alert=False)
            await query.edit_message_text(
                text=f"✅ Установлена фильтрация: {filter_names.get(filter_type, 'Неизвестно')}\n\n"
                     f"Режим: {'глобальный (sandbox)' if app_env == 'sandbox' else 'персональный (prod)'}"
            )
            return
        
        else:
            data = query.data or ""
            if ":" not in data:
                await query.answer("❌ Неизвестная команда", show_alert=False)
                return

            action, id_str = data.split(":", 1)
            if not id_str.isdigit():
                await query.answer("❌ Некорректный ID", show_alert=False)
                return

            news_id = int(id_str)
            user_id = query.from_user.id

            news = self.db.get_news_by_id(news_id) or self.news_cache.get(news_id)
            if not news:
                await query.answer("❌ Новость не найдена", show_alert=False)
                return

            category_tag = self._get_category_emoji(news.get('category', 'russia'))

            if action == "ai":
                try:
                    from config.config import AI_SUMMARY_MAX_REQUESTS_PER_MINUTE, APP_ENV
                    
                    # Check AI summary effective level
                    from core.services.access_control import get_effective_level
                    summary_level = get_effective_level(self.db, str(user_id), 'summary')
                    
                    if summary_level == 0:
                        await query.answer("⚠️ AI пересказ отключён администратором", show_alert=True)
                        return

                    now = time.time()
                    timestamps = self.user_ai_requests.get(user_id, [])
                    timestamps = [t for t in timestamps if now - t < 60]
                    if len(timestamps) >= AI_SUMMARY_MAX_REQUESTS_PER_MINUTE:
                        await query.answer("⏳ Слишком много запросов. Подождите минуту.", show_alert=False)
                        return
                    timestamps.append(now)
                    self.user_ai_requests[user_id] = timestamps

                    # Не отправляем отдельное сообщение, только query.answer без show_alert
                    # Это уменьшает количество сообщений и предотвращает flood control
                    await query.answer()
                    logger.info(f"AI summarize requested for news_id={news_id} by user={user_id}")

                    cached_summary = self.db.get_cached_summary(news_id)
                    if cached_summary:
                        # Check if already selected
                        is_selected = self.db.is_news_selected(user_id, news_id, env="prod")
                        select_btn_text = "✅ Выбрано" if is_selected else "📌 Выбрать"
                        
                        try:
                            await context.bot.send_message(
                                chat_id=user_id,
                                text=(
                                    f"🤖 Пересказ сгенерирован ИИ\n\n{cached_summary}\n\n"
                                    f"📰 Источник: {news.get('source', '')}\n{news.get('url', '')}"
                                ),
                                disable_web_page_preview=True,
                                disable_notification=True,
                                reply_markup=InlineKeyboardMarkup([[
                                    InlineKeyboardButton(select_btn_text, callback_data=f"select:{news_id}")
                                ]])
                            )
                        except RetryAfter as e:
                            # Flood control при отправке кешированного пересказа
                            await query.answer(
                                f"⏳ Слишком много сообщений.\nПопробуйте через {int(e.retry_after)} секунд.",
                                show_alert=True
                            )
                        return

                    lead_text = (
                        news.get('clean_text')
                        or news.get('lead_text')
                        or news.get('text', '')
                        or news.get('title', '')
                    )
                    from config.config import DEEPSEEK_INPUT_COST_PER_1K_TOKENS_USD, DEEPSEEK_OUTPUT_COST_PER_1K_TOKENS_USD

                    checksum = news.get('checksum')
                    logger.debug(f"Calling DeepSeek: lead_text_len={len(lead_text)}, title='{news.get('title', '')[:30]}', checksum={bool(checksum)}")
                    summary, token_usage = await self._summarize_with_deepseek(
                        lead_text,
                        news.get('title', ''),
                        checksum=checksum,
                        user_id=user_id,
                        bypass_tick_gate=True  # User requests bypass tick gate
                    )
                    logger.debug(f"DeepSeek response: summary={bool(summary)}, tokens={token_usage.get('total_tokens', 0)}")

                    if summary:
                        # Calculate cost based on input and output tokens
                        input_cost = (token_usage['input_tokens'] / 1000.0) * DEEPSEEK_INPUT_COST_PER_1K_TOKENS_USD
                        output_cost = (token_usage['output_tokens'] / 1000.0) * DEEPSEEK_OUTPUT_COST_PER_1K_TOKENS_USD
                        cost_usd = input_cost + output_cost
                        
                        self.db.add_ai_usage(tokens=token_usage['total_tokens'], cost_usd=cost_usd, operation_type='summarize')
                        self.db.save_summary(news_id, summary)
                        
                        # Check if already selected
                        is_selected = self.db.is_news_selected(user_id, news_id, env="prod")
                        select_btn_text = "✅ Выбрано" if is_selected else "📌 Выбрать"
                        
                        try:
                            await context.bot.send_message(
                                chat_id=user_id,
                                text=(
                                    f"🤖 Пересказ сгенерирован ИИ\n\n{summary}\n\n"
                                    f"📰 Источник: {news.get('source', '')}\n{news.get('url', '')}"
                                ),
                                disable_web_page_preview=True,
                                disable_notification=True,
                                reply_markup=InlineKeyboardMarkup([[
                                    InlineKeyboardButton(select_btn_text, callback_data=f"select:{news_id}")
                                ]])
                            )
                        except RetryAfter as e:
                            # Flood control при отправке нового пересказа
                            # Не бросаем исключение дальше, т.к. пересказ уже сохранён в кеш
                            logger.warning(f"Flood control when sending summary for {news_id}: retry after {e.retry_after}s")
                            await query.answer(
                                f"✅ Пересказ готов и сохранён!\n⏳ Показать через {int(e.retry_after)} сек (flood control)",
                                show_alert=True
                            )
                    else:
                        logger.warning(f"AI summarize failed for news_id={news_id}, no summary returned")
                        try:
                            await context.bot.send_message(
                                chat_id=user_id,
                                text="ИИ временно недоступен. Попробуйте позже.",
                                disable_web_page_preview=True,
                                disable_notification=True
                            )
                        except RetryAfter as e:
                            await query.answer(
                                f"⚠️ ИИ недоступен.\n⏳ Повторите через {int(e.retry_after)} сек",
                                show_alert=True
                            )
                    
                except RetryAfter as e:
                    # Telegram flood control - показываем сколько ждать
                    wait_seconds = int(e.retry_after)
                    logger.warning(f"Flood control for user {user_id}: retry after {wait_seconds}s")
                    try:
                        await query.answer(
                            f"⏳ Слишком много запросов.\nПопробуйте через {wait_seconds} секунд.",
                            show_alert=True
                        )
                    except:
                        pass
                except Exception as e:
                    logger.error(f"Error in AI summarize for news_id={news_id}: {e}", exc_info=True)
                    try:
                        # Используем query.answer вместо send_message чтобы избежать flood control
                        await query.answer(
                            f"❌ Ошибка: {str(e)[:80]}",
                            show_alert=True
                        )
                    except:
                        pass
                
                return
            
            elif action == "select":
                # Добавить/убрать новость из выбранных
                user_id = query.from_user.id
                
                if self.db.is_news_selected(user_id, news_id, env="prod"):
                    # Убрать из выбранных
                    self.db.remove_user_selection(user_id, news_id, env="prod")
                    await query.answer("✅ Убрано из выбранных", show_alert=False)
                    # Обновить кнопку
                    new_keyboard = InlineKeyboardMarkup([
                        [
                            InlineKeyboardButton("🤖 ИИ", callback_data=f"ai:{news_id}"),
                            InlineKeyboardButton("📌 Выбрать", callback_data=f"select:{news_id}")
                        ]
                    ])
                else:
                    # Добавить в выбранные
                    self.db.add_user_selection(user_id, news_id, env="prod")
                    await query.answer("✅ Добавлено в выбранные", show_alert=False)
                    # Обновить кнопку
                    new_keyboard = InlineKeyboardMarkup([
                        [
                            InlineKeyboardButton("🤖 ИИ", callback_data=f"ai:{news_id}"),
                            InlineKeyboardButton("✅ Выбрано", callback_data=f"select:{news_id}")
                        ]
                    ])
                
                # Обновить кнопки в сообщении
                try:
                    await query.edit_message_reply_markup(reply_markup=new_keyboard)
                except:
                    pass
                
                return

            elif action == "cluster":
                # Показать все источники в кластере
                cluster_id = int(parts[1])
                user_id = query.from_user.id
                
                try:
                    cluster_info = self.db.get_cluster_info(cluster_id)
                    if not cluster_info:
                        await query.answer("❌ Кластер не найден", show_alert=True)
                        return
                    
                    members = self.db.get_cluster_members(cluster_id)
                    if not members:
                        await query.answer("❌ Нет источников в кластере", show_alert=True)
                        return
                    
                    # Формируем сообщение со списком источников
                    message_lines = [
                        f"📰 Эта новость опубликована в {len(members)} источник{'ах' if len(members) > 4 else ('е' if len(members) == 1 else 'ах')}:\n"
                    ]
                    
                    for idx, member in enumerate(members, 1):
                        source_name = member.get('source', 'Unknown')
                        url = member.get('url', '')
                        # Truncate URL for display
                        display_url = url if len(url) < 50 else url[:47] + '...'
                        message_lines.append(f"{idx}. {source_name}\n   {display_url}")
                    
                    message_text = "\n".join(message_lines)
                    
                    await query.answer()
                    await context.bot.send_message(
                        chat_id=user_id,
                        text=message_text,
                        disable_web_page_preview=True,
                        disable_notification=True
                    )
                    return
                    
                except Exception as e:
                    logger.error(f"Error showing cluster {cluster_id}: {e}", exc_info=True)
                    await query.answer("❌ Ошибка при загрузке источников", show_alert=True)
                    return

            await query.answer("❌ Неизвестная команда", show_alert=False)
    
    async def _summarize_with_deepseek(self, text: str, title: str, checksum: str | None = None, user_id: int = None, bypass_tick_gate: bool = False) -> tuple[str | None, dict]:
        """
        Call DeepSeek API to summarize news.
        
        Args:
            text: Article text to summarize
            title: Article title
            user_id: User ID to get AI level preference (sandbox only)
            bypass_tick_gate: If True, skip tick gate check (for user-triggered requests)
            
        Returns:
            Tuple of (summary string or None, token usage dict)
        """
        try:
            from config.config import APP_ENV
            
            # Get effective AI level for summary
            from core.services.access_control import get_effective_level
            level = get_effective_level(self.db, str(user_id or 'global'), 'summary')

            # User-triggered requests bypass tick gate (they have their own rate limiting)
            if not bypass_tick_gate and not self._ai_tick_allow("summary"):
                return None, {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0, "skipped_by_gate": True}
            
            summary, token_usage = await self.deepseek_client.summarize(
                title=title,
                text=text,
                level=level,
                checksum=checksum
            )
            if summary:
                logger.debug(f"DeepSeek summary created (level={level}): {summary[:50]}...")
            return summary, token_usage
        except Exception as e:
            logger.error(f"DeepSeek error: {e}")
            return None, {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

    def _get_delivery_user_ids(self) -> list[int]:
        """Get user IDs to receive news in DM (approved users in prod, admins in sandbox)."""
        user_ids: set[int] = set()

        if get_app_env() == "prod":
            approved_users = self.access_db.get_approved_users()
            for user_id, _username, _first_name, _approved_at, _invited_by, _invite_label in approved_users:
                try:
                    user_ids.add(int(user_id))
                except Exception:
                    continue
        else:
            # Sandbox: send to admins only (no global user registry)
            user_ids.update(self.admin_ids)

        if not user_ids:
            # Fallback to admins if approved list is empty
            user_ids.update(self.admin_ids)

        return sorted(user_ids)

    async def _send_to_users(self, message: str, keyboard: InlineKeyboardMarkup, news_id: int, news_data: dict = None):
        """Отправляет новость пользователям в личные сообщения, учитывая их настройки источников и паузу"""
        recipients = self._get_delivery_user_ids()
        if not recipients:
            return

        # Prepare mapping source_code -> source_id once
        code_to_id = {}
        if news_data:
            sources = self.db.list_sources()
            code_to_id = {src['code']: src['id'] for src in sources}

        for user_id in recipients:
            await self._deliver_to_user(user_id, keyboard, news_id, news_data, message)

    async def _notify_realtime_no_news(self, interval_seconds: int) -> None:
        """Notify realtime users that there are no news items in the current tick."""
        if get_app_env() != "prod":
            return

        recipients = self._get_delivery_user_ids()
        if not recipients:
            return

        now_ts = time.time()
        message = (
            "🕒 Новостей пока нет.\n"
            f"Следующая проверка через {interval_seconds} сек."
        )

        for user_id in recipients:
            user_id_str = str(user_id)
            delivery_mode = self.db.get_user_delivery_mode(user_id_str, env="prod")
            if delivery_mode != "realtime":
                continue
            state = self.db.get_delivery_state(user_id_str, env="prod")
            if state.get("is_paused"):
                continue

            last_sent = self._last_no_news_sent_at.get(user_id)
            if last_sent and now_ts - last_sent < max(30, interval_seconds - 1):
                continue

            try:
                await self.application.bot.send_message(
                    chat_id=user_id,
                    text=message,
                    disable_web_page_preview=True,
                    disable_notification=True
                )
                self._last_no_news_sent_at[user_id] = now_ts
            except Exception as e:
                logger.warning(f"Failed to send no-news update to user {user_id}: {e}")

    async def _deliver_to_user(
        self,
        user_id: int,
        keyboard: InlineKeyboardMarkup,
        news_id: int,
        news_data: dict | None,
        fallback_message: str | None = None,
    ) -> bool:
        """Final delivery gate with pause/version checks and idempotency."""
        try:
            user_id_str = str(user_id)
            
            # Check delivery mode - if not realtime, buffer to pending digest
            delivery_mode = self.db.get_user_delivery_mode(user_id_str, env="prod")
            if delivery_mode in ('hourly', 'morning'):
                # Add to pending digest instead of immediate delivery
                buffered = self.db.add_to_pending_digest(user_id_str, news_id, delivery_mode)
                if buffered:
                    logger.info(f"DELIVERY_BUFFERED user_id={user_id} news_id={news_id} mode={delivery_mode}")
                return True
            
            if get_app_env() == "prod" and news_data:
                user_filter = self.db.get_user_category_filter(user_id_str, env="prod")
                if user_filter and news_data.get('category') != user_filter:
                    return False

            state_snapshot = self.db.get_delivery_state(user_id_str, env="prod")
            pause_version_snapshot = state_snapshot.get('pause_version', 0)

            if state_snapshot.get('is_paused'):
                logger.info(f"DELIVERY_SKIP_PAUSED user_id={user_id}")
                return False

            last_delivered = state_snapshot.get('last_delivered_news_id')
            if last_delivered and news_id <= int(last_delivered):
                logger.info(f"DELIVERY_DUPLICATE_SKIPPED user_id={user_id} news_id={news_id}")
                return False

            if news_data:
                enabled_source_ids = self.db.get_enabled_source_ids_for_user(user_id_str, env="prod")
                if enabled_source_ids is not None:
                    source = news_data.get('source', '')
                    source_id = None
                    if source:
                        source_id = None
                        if hasattr(self, '_source_code_to_id_cache') and self._source_code_to_id_cache:
                            source_id = self._source_code_to_id_cache.get(source)
                        else:
                            sources = self.db.list_sources()
                            self._source_code_to_id_cache = {src['code']: src['id'] for src in sources}
                            source_id = self._source_code_to_id_cache.get(source)
                    if source_id and source_id not in enabled_source_ids:
                        logger.info(f"DELIVERY_SKIP_SOURCE_DISABLED user_id={user_id} source={source}")
                        return False

            # Last gate: re-check pause/version
            state_current = self.db.get_delivery_state(user_id_str, env="prod")
            if state_current.get('is_paused'):
                logger.info(f"DELIVERY_SKIP_PAUSED user_id={user_id}")
                return False
            if state_current.get('pause_version', 0) != pause_version_snapshot:
                logger.info(f"DELIVERY_SKIP_VERSION_MISMATCH user_id={user_id}")
                return False

            if not self.db.try_log_delivery(user_id_str, news_id):
                logger.info(f"DELIVERY_DUPLICATE_SKIPPED user_id={user_id} news_id={news_id}")
                return False

            message_to_send = fallback_message or ''
            if news_data:
                base_text = (
                    news_data.get('clean_text')
                    or news_data.get('text', '')
                    or news_data.get('lead_text', '')
                )
                title = news_data.get('title', 'No title')
                source_name = news_data.get('source', 'Unknown')
                source_url = news_data.get('url', '')
                translate_enabled, target_lang = self.db.get_user_translation(user_id_str, env="prod")
                translated_text = None
                if translate_enabled and news_data.get('language') == 'en' and base_text:
                    checksum = news_data.get('checksum') or ''
                    if checksum:
                        translated_text = self.db.get_translation_cache(news_id, checksum, target_lang)
                    if not translated_text:
                        translated_text, token_usage = await self.deepseek_client.translate_text(
                            base_text,
                            target_lang=target_lang,
                            checksum=checksum or None
                        )
                        if translated_text and checksum:
                            self.db.set_translation_cache(news_id, checksum, target_lang, translated_text)
                        if token_usage and token_usage.get('total_tokens', 0) > 0:
                            cost_usd = token_usage.get('cost_usd', 0.0) or 0.0
                            self.db.add_ai_usage(token_usage['total_tokens'], cost_usd, 'translate')

                language = news_data.get('language') or 'ru'
                tag_language = 'ru' if (translate_enabled and language == 'en') else ('en' if language == 'en' else 'ru')
                extra_tags = news_data.get('hashtags_ru') if tag_language == 'ru' else news_data.get('hashtags_en')
                extra_tags = extra_tags or ''

                message_to_send = format_telegram_message(
                    title=title,
                    text=translated_text or base_text,
                    source_name=source_name,
                    source_url=source_url,
                    category=self._get_category_line(
                        news_data.get('category', 'russia'),
                        language=tag_language,
                        extra_tags=extra_tags
                    )
                )

            await self.application.bot.send_message(
                chat_id=user_id,
                text=message_to_send,
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=keyboard,
                disable_web_page_preview=True,
                disable_notification=True
            )
            self.db.update_last_delivered(user_id_str, news_id, env="prod")
            logger.info(f"DELIVERY_SENT_OK user_id={user_id} news_id={news_id}")
            return True
        except Exception as e:
            logger.warning(f"Failed to send to user {user_id}: {e}")
            self.db.remove_delivery_log(str(user_id), news_id)
            return False

    async def _deliver_pending_for_user(self, user_id: int, limit: int = 50):
        """Deliver pending news to user after resume."""
        try:
            state = self.db.get_delivery_state(str(user_id), env="prod")
            last_id = state.get('last_delivered_news_id')
            pending = self.db.get_news_after_id(last_id, limit=limit)
            if not pending:
                return

            user_filter = None
            if get_app_env() == "prod":
                user_filter = self.db.get_user_category_filter(str(user_id), env="prod")

            for item in pending:
                if not self._is_today_news(item):
                    continue
                if user_filter and item.get('category') != user_filter:
                    continue
                news_id = item.get('id')
                if not news_id:
                    continue

                full = self.db.get_news_by_id(int(news_id)) or item
                keyboard = InlineKeyboardMarkup([
                    [
                        InlineKeyboardButton("🤖 ИИ", callback_data=f"ai:{news_id}"),
                        InlineKeyboardButton("📌 Выбрать", callback_data=f"select:{news_id}")
                    ]
                ])

                await self._deliver_to_user(user_id, keyboard, int(news_id), full, None)
        except Exception as e:
            logger.error(f"Error delivering pending for user {user_id}: {e}")
    
    async def collect_and_publish(self, force: bool = False) -> int:
        """
        Собирает новости и публикует их
        Возвращает количество опубликованных новостей
        """
        from core.services.global_stop import get_global_stop
        
        # Проверяем глобальный стоп (постоянный, для всей системы)
        if get_global_stop():
            logger.info("Global system stop is ON, skipping collection")
            return 0
        
        # Global collection stop flag (hard stop for prod + sandbox)
        stop_state = get_global_collection_stop_state(app_env=get_app_env())
        if stop_state.enabled:
            logger.info(
                f"tick_skipped_global_stop ttl_sec_remaining={stop_state.ttl_sec_remaining} key={stop_state.key}"
            )
            return 0
        
        if self.is_paused:
            logger.info("Bot is paused, skipping collection")
            return 0
        
        # Prevent concurrent collection cycles
        if self.collection_lock.locked():
            logger.info("Collection already in progress, skipping")
            return 0
        
        async with self.collection_lock:
            return await self._do_collect_and_publish(force=force)
    
    async def run_tier_adjustment(self):
        """Periodic task to auto-adjust source tiers based on quality score."""
        logger.info("Starting tier adjustment scheduler (checks daily)")
        while True:
            try:
                await asyncio.sleep(24 * 60 * 60)  # Run once per day
                
                logger.info("Running auto-adjustment of source tiers...")
                result = self.db.auto_adjust_source_tiers(days=7, promote_threshold=0.8, demote_threshold=0.6)
                
                if result['promoted']:
                    logger.info(f"Promoted sources: {result['promoted']}")
                if result['demoted']:
                    logger.info(f"Demoted sources: {result['demoted']}")
                
                logger.info(f"Tier adjustment complete: {len(result['promoted'])} promoted, {len(result['demoted'])} demoted")
            except Exception as e:
                logger.error(f"Error in tier adjustment: {e}", exc_info=True)
                await asyncio.sleep(60 * 60)  # Wait 1 hour on error before retry

    async def run_hourly_digest(self):
        """Periodic task to send hourly digests."""
        logger.info("Starting hourly digest scheduler")
        # Initial delay to align with hour boundary
        import datetime
        now = datetime.datetime.now()
        minutes_until_next_hour = 60 - now.minute
        await asyncio.sleep(minutes_until_next_hour * 60)
        
        while True:
            try:
                logger.info("Running hourly digest delivery...")
                await self._send_digest_to_users('hourly')
            except Exception as e:
                logger.error(f"Error in hourly digest: {e}", exc_info=True)
            
            await asyncio.sleep(60 * 60)  # Run every hour

    async def run_morning_digest(self):
        """Periodic task to send morning digests at 7:00 AM."""
        logger.info("Starting morning digest scheduler")
        
        while True:
            try:
                import datetime
                now = datetime.datetime.now()
                target_hour = 7  # 7:00 AM
                
                # Calculate seconds until next 7:00 AM
                target = now.replace(hour=target_hour, minute=0, second=0, microsecond=0)
                if now >= target:
                    target += datetime.timedelta(days=1)
                
                wait_seconds = (target - now).total_seconds()
                logger.info(f"Next morning digest in {wait_seconds/3600:.1f} hours")
                await asyncio.sleep(wait_seconds)
                
                logger.info("Running morning digest delivery...")
                await self._send_digest_to_users('morning')
            except Exception as e:
                logger.error(f"Error in morning digest: {e}", exc_info=True)
                await asyncio.sleep(60 * 60)  # Wait 1 hour on error

    async def _send_digest_to_users(self, delivery_mode: str):
        """Send pending digest items to all users with the specified delivery mode."""
        try:
            user_ids = self.db.get_users_by_delivery_mode(delivery_mode, env="prod")
            if not user_ids:
                logger.info(f"No users with {delivery_mode} delivery mode")
                return
            
            logger.info(f"Sending {delivery_mode} digest to {len(user_ids)} users")
            sent_count = 0
            
            for user_id_str in user_ids:
                try:
                    pending = self.db.get_pending_digest_items(user_id_str, delivery_mode)
                    if not pending:
                        continue
                    
                    # Send digest header
                    if delivery_mode == 'hourly':
                        header = f"📰 Почасовая подборка ({len(pending)} новостей)"
                    else:
                        header = f"☀️ Утренний дайджест ({len(pending)} новостей)"
                    
                    user_id = int(user_id_str)
                    await self.application.bot.send_message(
                        chat_id=user_id,
                        text=header,
                        disable_notification=True
                    )
                    
                    # Send each news item
                    news_ids_sent = []
                    for item in pending[:50]:  # Limit to 50 items per digest
                        news_id = item['news_id']
                        news_data = self.db.get_news_by_id(news_id)
                        if not news_data:
                            continue
                        
                        # Check if news passes user filters
                        user_filter = self.db.get_user_category_filter(user_id_str, env="prod")
                        if user_filter and news_data.get('category') != user_filter:
                            continue
                        
                        enabled_source_ids = self.db.get_enabled_source_ids_for_user(user_id_str, env="prod")
                        if enabled_source_ids is not None:
                            source = news_data.get('source', '')
                            if source:
                                if not hasattr(self, '_source_code_to_id_cache'):
                                    sources = self.db.list_sources()
                                    self._source_code_to_id_cache = {src['code']: src['id'] for src in sources}
                                source_id = self._source_code_to_id_cache.get(source)
                                if source_id and source_id not in enabled_source_ids:
                                    continue
                        
                        # Build message
                        base_text = (
                            news_data.get('clean_text')
                            or news_data.get('text', '')
                            or news_data.get('lead_text', '')
                        )
                        title = news_data.get('title', 'No title')
                        source_name = news_data.get('source', 'Unknown')
                        source_url = news_data.get('url', '')
                        
                        # Translation if enabled
                        translate_enabled, target_lang = self.db.get_user_translation(user_id_str, env="prod")
                        translated_text = None
                        if translate_enabled and news_data.get('language') == 'en' and base_text:
                            checksum = news_data.get('checksum') or ''
                            if checksum:
                                translated_text = self.db.get_translation_cache(news_id, checksum, target_lang)
                        
                        language = news_data.get('language') or 'ru'
                        tag_language = 'ru' if (translate_enabled and language == 'en') else ('en' if language == 'en' else 'ru')
                        extra_tags = news_data.get('hashtags_ru') if tag_language == 'ru' else news_data.get('hashtags_en')
                        extra_tags = extra_tags or ''
                        
                        message = format_telegram_message(
                            title=title,
                            text=translated_text or base_text,
                            source_name=source_name,
                            source_url=source_url,
                            category=self._get_category_line(
                                news_data.get('category', 'russia'),
                                language=tag_language,
                                extra_tags=extra_tags
                            )
                        )
                        
                        # Build keyboard with cluster button if applicable
                        cluster_id = self.db.get_cluster_for_news(news_id)
                        cluster_info = None
                        if cluster_id:
                            cluster_info = self.db.get_cluster_info(cluster_id)
                        
                        buttons_row1 = [
                            InlineKeyboardButton("✨ ИИ", callback_data=f"ai:{news_id}"),
                            InlineKeyboardButton("✅ Отбор", callback_data=f"select:{news_id}")
                        ]
                        buttons_rows = [buttons_row1]
                        
                        if cluster_info and cluster_info['member_count'] > 1:
                            source_count = cluster_info['member_count']
                            suffix = "ов" if source_count > 4 else ("а" if source_count in (2, 3, 4) else "")
                            buttons_rows.append([
                                InlineKeyboardButton(
                                    f"📰 +{source_count-1} источник{suffix}",
                                    callback_data=f"cluster:{cluster_id}"
                                )
                            ])
                        
                        keyboard = InlineKeyboardMarkup(buttons_rows)
                        
                        await self.application.bot.send_message(
                            chat_id=user_id,
                            text=message,
                            parse_mode=ParseMode.MARKDOWN,
                            reply_markup=keyboard,
                            disable_web_page_preview=True,
                            disable_notification=True
                        )
                        
                        news_ids_sent.append(news_id)
                        self.db.update_last_delivered(user_id_str, news_id, env="prod")
                    
                    # Clear sent items from pending
                    if news_ids_sent:
                        cleared = self.db.clear_pending_digest(user_id_str, news_ids_sent)
                        logger.info(f"Sent {len(news_ids_sent)} digest items to user {user_id}, cleared {cleared}")
                        sent_count += len(news_ids_sent)
                
                except Exception as e:
                    logger.error(f"Error sending {delivery_mode} digest to user {user_id_str}: {e}")
            
            logger.info(f"Completed {delivery_mode} digest: sent {sent_count} items total")
        except Exception as e:
            logger.error(f"Error in _send_digest_to_users: {e}", exc_info=True)
    
    async def _do_collect_and_publish(self, force: bool = False) -> int:
        """
        Internal method: performs the actual collection and publishing
        """
        try:
            # Собираем новости
            logger.info("Starting news collection...")
            stop_state = get_global_collection_stop_state(app_env=get_app_env())
            if stop_state.enabled:
                logger.info(
                    {
                        "event": "collect_skipped_global_stop",
                        "ttl_sec_remaining": stop_state.ttl_sec_remaining,
                        "key": stop_state.key,
                    }
                )
                return 0
            tick_id = datetime.utcnow().strftime("%Y%m%dT%H%M%S")
            self._begin_ai_tick(tick_id)
            cache_hits_start = 0
            if self.deepseek_client.cache:
                try:
                    cache_hits_start = int(self.deepseek_client.cache.get_stats().get("hits", 0) or 0)
                except Exception:
                    cache_hits_start = 0
            news_items = await self.collector.collect_all(force=force)
            
            # Send admin notifications for quarantined sources
            if self.collector.quarantined_sources_this_tick:
                await self._notify_admins_quarantine(self.collector.quarantined_sources_this_tick)

            app_env = get_app_env()
            global_category_filter = self._get_global_category_filter() if app_env == "sandbox" else None
            
            published_count = 0
            max_publications = 100  # Лимит публикаций за цикл (защита от rate limiting)
            
            # Track per-source statistics for quality metrics
            source_stats = {}  # {source: {'total': int, 'new': int, 'duplicate': int}}
            
            # Кэш дубликатов в текущей сессии (защита от повторов в одном цикле)
            session_titles = set()  # normalized titles for duplicate detection
            session_url_hashes = set()
            session_checksums = set()
            session_url_normalized = set()
            recent_simhashes = self.db.get_recent_simhashes(hours=48, limit=1500)
            
            # Публикуем каждую новость
            for news in news_items:
                # Track per-source statistics
                source = news.get('source', '')
                if source:
                    if source not in source_stats:
                        source_stats[source] = {'total': 0, 'new': 0, 'duplicate': 0}
                    source_stats[source]['total'] += 1
                
                # Stop may be toggled while processing a tick.
                if get_global_collection_stop_state(app_env=get_app_env()).enabled:
                    logger.info({"event": "publish_aborted_global_stop"})
                    break
                # Ensure fetched_at and URL fingerprints are present
                if not news.get('fetched_at'):
                    news['fetched_at'] = datetime.utcnow().isoformat()
                if not news.get('url_normalized') and news.get('url'):
                    news['url_normalized'] = normalize_url(news.get('url'))
                if not news.get('url_hash') and news.get('url'):
                    url_for_hash = news.get('url_normalized') or news.get('url')
                    news['url_hash'] = compute_url_hash(url_for_hash)
                if not news.get('simhash'):
                    title_for_hash = news.get('title', '')
                    text_for_hash = news.get('clean_text') or news.get('text', '')
                    news['simhash'] = compute_simhash(text_for_hash, title=title_for_hash)

                # First-seen check for confidence=none
                if (news.get('published_confidence') or 'none').lower() == 'none':
                    guid = news.get('guid')
                    url_hash = news.get('url_hash')
                    news['is_first_seen'] = not self.db.is_seen_guid_or_url_hash(guid, url_hash)
                    if news.get('is_first_seen') and not news.get('first_seen_at'):
                        news['first_seen_at'] = news.get('fetched_at')

                ok, reason = self._should_publish_news(news)
                if not ok:
                    domain = self._get_domain(news)
                    self._record_drop_reason(domain, reason)
                    source = news.get('source', '')
                    if reason == "OLD_PUBLISHED_AT":
                        self.db.record_source_event(source, "drop_old")
                    elif reason in ("NO_PUBLISHED_DATE", "PARSE_DATE_FAILED"):
                        self.db.record_source_event(source, "drop_date", error_code=reason)
                    else:
                        self.db.record_source_event(source, "error", error_code=reason)
                    logger.debug(f"Skipping news ({reason}): {news.get('title', '')[:50]}")
                    continue

                # Проверяем лимит публикаций
                if published_count >= max_publications:
                    logger.info(f"Reached publication limit ({max_publications}), stopping")
                    break
                
                # Проверяем фильтр по источникам для пользователя (система admin_ids)
                # TELEGRAM_CHANNEL_ID - основной канал, где видят все подписчики
                # Но админы в ADMIN_IDS могут видеть разные выборки
                # На данный момент - выдача всем одинаковая (глобальная)
                
                # Проверяем фильтр по категориям (sandbox global)
                if global_category_filter and news.get('category') != global_category_filter:
                    logger.debug(f"Skipping news (category filter): {news.get('title')[:50]}")
                    continue
                
                # Проверяем дубликат в текущей сессии (быстрая проверка)
                import re
                title = news.get('title', '')
                normalized = re.sub(r'[^\w\s]', '', title.lower())
                if normalized in session_titles:
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate in session: {title[:50]}")
                    continue
                session_titles.add(normalized)

                url_hash = news.get('url_hash') or ''
                if url_hash and url_hash in session_url_hashes:
                    logger.debug(f"Skipping duplicate url_hash in session: {title[:50]}")
                    continue
                if url_hash:
                    session_url_hashes.add(url_hash)

                checksum = news.get('checksum') or ''
                if checksum and checksum in session_checksums:
                    logger.debug(f"Skipping duplicate checksum in session: {title[:50]}")
                    continue
                if checksum:
                    session_checksums.add(checksum)

                url_normalized = news.get('url_normalized') or ''
                if url_normalized and url_normalized in session_url_normalized:
                    logger.debug(f"Skipping duplicate url_normalized in session: {title[:50]}")
                    continue
                if url_normalized:
                    session_url_normalized.add(url_normalized)

                # Проверка дубликатов по URL hash / guid / URL canonical
                if self.db.is_seen_guid_or_url_hash(news.get('guid'), url_hash):
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate guid/url_hash: {title[:50]}")
                    continue
                if url_normalized and self.db.is_url_normalized_seen(url_normalized):
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate url_normalized: {title[:50]}")
                    continue

                # Проверка дубликатов по checksum (контент) в окне 48 часов
                if checksum and self.db.is_checksum_recent(checksum, hours=48):
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate checksum: {title[:50]}")
                    continue

                # Проверка дубликатов по content_hash (нормализованный title+text) в окне 48 часов
                content_hash = news.get('content_hash') or ''
                if content_hash and self.db.is_content_hash_recent(content_hash, hours=48):
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate content_hash: {title[:50]}")
                    continue

                # Проверка near-duplicate по simhash
                simhash = news.get('simhash')
                if isinstance(simhash, int) and recent_simhashes:
                    for existing in recent_simhashes:
                        if hamming_distance(simhash, existing) <= 6:
                            logger.debug(f"Skipping near-duplicate simhash: {title[:50]}")
                            simhash = None
                            break
                if simhash is None and news.get('simhash') is not None:
                    continue
                
                # Проверяем дубликат по заголовку в БД (защита от одной новости на разных источниках)
                if self.db.is_similar_title_published(title, threshold=0.85):  # Increased threshold to 0.85
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping similar title: {title[:50]}")
                    continue
                
                # Попытка атомарно зарегистрировать новость в БД
                hashtags_ru = ""
                hashtags_en = ""
                try:
                    hashtags_ru, hashtags_en = await self._generate_hashtags_snapshot(news)
                except Exception as e:
                    logger.debug(f"Hashtags generation skipped: {e}")

                news_id = self.db.add_news(
                    url=news['url'],
                    title=news.get('title', ''),
                    source=news.get('source', ''),
                    category=news.get('category', ''),
                    lead_text=news.get('lead_text', '') or news.get('text', '') or '',
                    raw_text=news.get('raw_text'),
                    clean_text=news.get('clean_text') or news.get('text', ''),
                    checksum=news.get('checksum'),
                    content_hash=news.get('content_hash'),
                    language=news.get('language'),
                    domain=news.get('domain'),
                    extraction_method=news.get('extraction_method'),
                    published_at=news.get('published_at'),
                    published_date=news.get('published_date'),
                    published_time=news.get('published_time'),
                    published_confidence=news.get('published_confidence'),
                    published_source=news.get('published_source'),
                    fetched_at=news.get('fetched_at'),
                    first_seen_at=news.get('first_seen_at') or news.get('fetched_at'),
                    url_hash=news.get('url_hash'),
                    url_normalized=news.get('url_normalized'),
                    guid=news.get('guid'),
                    simhash=news.get('simhash'),
                    quality_score=news.get('quality_score'),
                    hashtags_ru=hashtags_ru,
                    hashtags_en=hashtags_en,
                )

                if not news_id:
                    if source:
                        source_stats[source]['duplicate'] += 1
                    logger.debug(f"Skipping duplicate URL: {news.get('url')}")
                    continue

                # Successfully added new item
                if source:
                    source_stats[source]['new'] += 1

                if isinstance(news.get('simhash'), int):
                    recent_simhashes.insert(0, news['simhash'])

                self.db.record_source_event(news.get('source', ''), "success")

                # Event clustering: group similar news from different sources
                cluster_id = None
                if isinstance(news.get('simhash'), int):
                    try:
                        # Find similar clusters within 6-hour window (tighter threshold for clustering)
                        similar_clusters = self.db.find_similar_clusters(
                            news['simhash'], 
                            hours=6, 
                            hamming_threshold=3
                        )
                        
                        if similar_clusters:
                            # Add to existing cluster (use first match)
                            cluster_id = similar_clusters[0]
                            self.db.add_news_to_cluster(cluster_id, news_id)
                            cluster_info = self.db.get_cluster_info(cluster_id)
                            if cluster_info:
                                logger.info(
                                    f"Added news {news_id} to cluster {cluster_id} "
                                    f"(now {cluster_info['member_count']} sources)"
                                )
                        else:
                            # Create new cluster with this news as representative
                            cluster_id = self.db.create_cluster(news_id)
                            if cluster_id:
                                logger.debug(f"Created new cluster {cluster_id} for news {news_id}")
                    except Exception as e:
                        logger.debug(f"Error in event clustering: {e}")

                # Check if we need auto-summarization for all sources (cleanup_level=5)
                from core.services.access_control import AILevelManager
                ai_manager = AILevelManager(self.db)
                cleanup_level = ai_manager.get_level('global', 'cleanup')
                
                source = news.get('source', '').lower()
                news_text = news.get('clean_text') or news.get('text', '')
                
                # Debug logging for auto-summarization trigger
                logger.debug(f"Auto-summarize check: cleanup_level={cleanup_level}, source={source}")
                
                # Auto-summarize ALL sources when cleanup_level=5
                if cleanup_level == 5:
                    logger.info(f"Auto-summarizing {source} (cleanup_level=5)")
                    try:
                        # Get or generate summary
                        cached_summary = self.db.get_cached_summary(news_id)
                        if cached_summary:
                            logger.debug(f"Using cached summary for {news_id}")
                            news_text = cached_summary
                        else:
                            # Generate summary (1-2 sentences)
                            full_text = news_text if news_text else news.get('title', '')
                            summary_level = ai_manager.get_level('global', 'summary')
                            checksum = news.get('checksum')

                            summary = None
                            if self._ai_tick_allow("summary"):
                                summary, _usage = await self.deepseek_client.summarize(
                                    title=news.get('title', ''),
                                    text=full_text[:2000],
                                    level=summary_level,
                                    checksum=checksum
                                )
                            else:
                                logger.info("AI summary skipped by tick gate")

                            if summary:
                                self.db.save_summary(news_id, summary)
                                news_text = summary
                                logger.info(f"Generated auto-summary for {source}: {summary[:50]}...")
                            else:
                                logger.warning(f"Summarization returned empty result for {source}")
                    except Exception as e:
                        logger.error(f"Error auto-summarizing {source}: {e}", exc_info=True)
                
                # Формируем сообщение
                news_category = news.get('category', 'russia')
                category_emoji = self._get_category_emoji(news_category)
                
                # Debug: логируем текст перед форматированием
                text_preview = news_text[:100] if news_text else "(no text)"
                logger.debug(f"Formatting message: title={news.get('title', '')[:40]}... text={text_preview}...")
                
                message = format_telegram_message(
                    title=news.get('title', 'No title'),
                    text=news_text,
                    source_name=news.get('source', 'Unknown'),
                    source_url=news.get('url', ''),
                    category=category_emoji
                )
                
                # Сохраняем в кэш для ИИ кнопки
                self.news_cache[news_id] = {
                    'title': news.get('title', 'No title'),
                    'text': news_text,
                    'lead_text': news_text,
                    'url': news.get('url', ''),
                    'source': news.get('source', 'Unknown'),
                    'category': news_category,
                    'clean_text': news.get('clean_text') or news_text,
                    'checksum': news.get('checksum'),
                    'url_normalized': news.get('url_normalized'),
                    'simhash': news.get('simhash'),
                    'language': news.get('language'),
                    'published_date': news.get('published_date'),
                    'published_time': news.get('published_time'),
                    'hashtags_ru': hashtags_ru,
                    'hashtags_en': hashtags_en,
                }

                # Получаем информацию о кластере (если новость в кластере)
                cluster_info = None
                if cluster_id:
                    cluster_info = self.db.get_cluster_info(cluster_id)

                # Создаем кнопки: ИИ пересказ, Выбрать, и опционально Источники
                buttons_row1 = [
                    InlineKeyboardButton("🤖 ИИ", callback_data=f"ai:{news_id}"),
                    InlineKeyboardButton("📌 Выбрать", callback_data=f"select:{news_id}")
                ]
                
                # Если в кластере больше 1 источника, добавляем кнопку "Источники"
                buttons_rows = [buttons_row1]
                if cluster_info and cluster_info['member_count'] > 1:
                    source_count = cluster_info['member_count']
                    buttons_rows.append([
                        InlineKeyboardButton(
                            f"📰 +{source_count - 1} источник{'ов' if source_count > 4 else ('а' if source_count <= 3 else 'ов')}", 
                            callback_data=f"cluster:{cluster_id}"
                        )
                    ])
                
                keyboard = InlineKeyboardMarkup(buttons_rows)

                try:
                    logger.info(f"Channel publish skipped (prod bot only): {news['title'][:50]}")
                    
                    # Сохраняем news_id как опубликованную (для корректной статистики)
                    published_count += 1
                    
                    # Отправляем пользователям в личку с кнопкой "ИИ" и учётом их настроек источников
                    await self._send_to_users(message, keyboard, news_id, news)

                    # Задержка между публикациями (защита от Telegram rate limiting)
                    await asyncio.sleep(0.5)  # Меньше задержка так как не отправляем в канал

                except Exception as e:
                    logger.error(f"Error publishing news: {type(e).__name__} (URL hidden)")
                    # Откатываем запись в БД, чтобы можно было попытаться снова
                    try:
                        self.db.remove_news_by_url(news['url'])
                    except Exception:
                        pass
            
            logger.info(f"Collection complete. Published {published_count} new items")
            if self.drop_counters:
                logger.info(f"Drop reasons summary: {self.drop_counters}")
                self.drop_counters = {}

            cache_hits_end = cache_hits_start
            if self.deepseek_client.cache:
                try:
                    cache_hits_end = int(self.deepseek_client.cache.get_stats().get("hits", 0) or 0)
                except Exception:
                    cache_hits_end = cache_hits_start
            cache_hits_tick = max(0, cache_hits_end - cache_hits_start)

            budget_state = "OK"
            if self.deepseek_client.budget:
                try:
                    budget_state = self.deepseek_client.budget.get_state().get("budget_state", "OK")
                except Exception:
                    budget_state = "OK"

            tick_state = self._get_ai_tick_state()
            tick_log = {
                "tick_id": tick_id,
                "fetched": len(news_items),
                "parsed": len(news_items),
                "deduped": max(0, len(news_items) - published_count),
                "published": published_count,
                "ai_calls": tick_state.get("calls", 0),
                "ai_cache_hits": cache_hits_tick,
                "budget_state": budget_state,
            }
            
            # Record source quality statistics
            for source_name, stats in source_stats.items():
                if stats['total'] > 0:  # Only record if source had items
                    self.db.update_source_quality_stats(
                        source_name,
                        stats['total'],
                        stats['new'],
                        stats['duplicate']
                    )
            
            logger.info("TICK_STATS %s", json.dumps(tick_log, ensure_ascii=True))

            if published_count == 0:
                await self._notify_realtime_no_news(CHECK_INTERVAL_SECONDS)
            return published_count
        
        except Exception as e:
            logger.error(f"Error in collect_and_publish: {e}")
            return 0
    
    async def _notify_admins_quarantine(self, quarantined: list):
        """Send notifications to admins about quarantined sources."""
        if not quarantined:
            return
        
        try:
            # Build notification message
            message_lines = ["🔴 ИСТОЧНИКИ В КАРАНТИНЕ\n"]
            for info in quarantined:
                source = info.get('source', 'Unknown')
                reason = info.get('reason', 'unknown')
                error_streak = info.get('error_streak', 0)
                error_code = info.get('last_error_code', 'N/A')
                
                message_lines.append(
                    f"📰 {source}\n"
                    f"   Причина: {reason}\n"
                    f"   Ошибок подряд: {error_streak}\n"
                    f"   Код ошибки: {error_code}\n"
                )
            
            message_lines.append("\nИсточники автоматически отключены. Используйте /admin → Источники для восстановления.")
            message = "\n".join(message_lines)
            
            # Send to all admins
            for admin_id in self.admin_ids:
                try:
                    await self.application.bot.send_message(
                        chat_id=admin_id,
                        text=message,
                        disable_notification=False  # Important notification
                    )
                except Exception as e:
                    logger.error(f"Failed to notify admin {admin_id} about quarantine: {e}")
            
            logger.info(f"Sent quarantine notifications for {len(quarantined)} sources to {len(self.admin_ids)} admins")
        
        except Exception as e:
            logger.error(f"Error sending quarantine notifications: {e}")
    
    def _get_category_emoji(self, category: str) -> str:
        """Возвращает категорию с эмодзи и хештегом"""
        from config.config import CATEGORIES
        return CATEGORIES.get(category, 'Новости')

    def _get_category_tag(self, category: str, language: str = 'ru') -> str:
        """Return category hashtag string for RU/EN."""
        if language == 'en':
            mapping = {
                'world': '#World',
                'russia': '#Russia',
                'moscow': '#Moscow',
                'moscow_region': '#MoscowRegion',
            }
            return mapping.get(category, '')
        mapping = {
            'world': '#Мир',
            'russia': '#Россия',
            'moscow': '#Москва',
            'moscow_region': '#Подмосковье',
        }
        return mapping.get(category, '')

    def _normalize_hashtag(self, tag: str) -> str:
        cleaned = (tag or '').strip()
        if not cleaned:
            return ''
        if not cleaned.startswith('#'):
            cleaned = '#' + cleaned
        return cleaned

    def _enforce_category_hashtag(self, tags: list[str], category: str, language: str) -> list[str]:
        required = self._get_category_tag(category, language)
        normalized = []
        seen = set()
        for tag in tags:
            cleaned = self._normalize_hashtag(tag)
            if not cleaned:
                continue
            key = cleaned.lower()
            if key not in seen:
                normalized.append(cleaned)
                seen.add(key)

        required_key = required.lower()
        if required_key not in seen:
            normalized.insert(0, required)
            seen.add(required_key)

        if category in ('moscow', 'moscow_region'):
            drop = '#россия' if language == 'ru' else '#russia'
            normalized = [tag for tag in normalized if tag.lower() != drop]

        return normalized[:8]

    def _get_category_line(self, category: str, language: str = 'ru', extra_tags: str = '') -> str:
        """Return category line with emoji and optional extra hashtags."""
        emoji_map = {
            'world': '🌍',
            'russia': '🇷🇺',
            'moscow': '🏛️',
            'moscow_region': '🏘️',
        }
        emoji = emoji_map.get(category, '🗞')
        base_tag = self._get_category_tag(category, language)
        tags = extra_tags.strip() if extra_tags else base_tag
        return f"{emoji} {tags}".strip()

    async def _generate_hashtags_snapshot(self, news: dict) -> tuple[str, str]:
        """Generate and return (hashtags_ru, hashtags_en) strings."""
        title = news.get('title', '')
        text = news.get('clean_text') or news.get('text', '') or ''
        language = news.get('language') or 'ru'
        category = news.get('category', 'russia')
        from core.services.access_control import get_effective_level
        from utils.hashtags import build_hashtags, build_hashtags_en

        level = get_effective_level(self.db, 'global', 'hashtags')

        try:
            tags_ru = await build_hashtags(
                title=title,
                lead_text=text,
                source='global',
                existing_category=category,
                language=language,
                ai_client=self.deepseek_client,
                level=level,
                ai_call_guard=self._ai_tick_allow,
            )
        except Exception as e:
            logger.debug(f"Hashtag taxonomy failed: {e}")
            tags_ru = []

        if not tags_ru:
            tags_ru = ["#Россия", "#Общество"]

        tags_en = build_hashtags_en(tags_ru)

        hashtags_ru = " ".join(tags_ru)
        hashtags_en = " ".join(tags_en)
        return hashtags_ru, hashtags_en
    
    async def run_periodic_collection(self):
        """Запускает периодический сбор новостей"""
        from core.services.global_stop import get_global_stop, wait_global_stop, wait_for_resume
        
        logger.info("Starting periodic news collection")
        
        while self.is_running:
            try:
                # Проверяем глобальный стоп
                if get_global_stop():
                    logger.info("⏸️  Global stop activated - waiting for resume signal...")
                    await wait_for_resume()
                    logger.info("⏯️  Resuming collection after signal")
                    continue

                if self.is_paused:
                    try:
                        await asyncio.wait_for(wait_global_stop(), timeout=CHECK_INTERVAL_SECONDS)
                    except asyncio.TimeoutError:
                        pass
                    continue

                collection_task = asyncio.create_task(self.collect_and_publish())
                stop_task = asyncio.create_task(wait_global_stop())
                done, pending = await asyncio.wait(
                    {collection_task, stop_task},
                    return_when=asyncio.FIRST_COMPLETED,
                )

                if stop_task in done:
                    logger.info("⏹️  Global stop during collection - cancelling in-flight work")
                    collection_task.cancel()
                    with suppress(asyncio.CancelledError):
                        await collection_task
                    await wait_for_resume()
                    continue

                stop_task.cancel()
                with suppress(asyncio.CancelledError):
                    await stop_task

                # Ждем перед следующей проверкой
                # (Если стоп активируется, wait_global_stop() сработает мгновенно)
                try:
                    await asyncio.wait_for(wait_global_stop(), timeout=CHECK_INTERVAL_SECONDS)
                except asyncio.TimeoutError:
                    pass
            
            except Exception as e:
                logger.error(f"Error in periodic collection: {e}")
                await asyncio.sleep(CHECK_INTERVAL_SECONDS)
    
    async def start(self):
        """Запускает бота"""
        logger.info("Starting bot...")

        try:
            from config.config import APP_ENV
            if APP_ENV == "sandbox":
                self.db.reset_bot_lock()
        except Exception:
            pass

        if not self._acquire_instance_lock():
            return

        if not self.db.acquire_bot_lock(self._db_instance_id, ttl_seconds=600):
            self._release_instance_lock()
            return
        
        # Инициализируем админов в БД (при первом запуске)
        self._init_admins_access()
        
        # Инициализируем asyncio.Event для глобального стопа (мгновенная отмена задач)
        from core.services.global_stop import init_global_stop_event
        await init_global_stop_event()
        logger.info("Global stop event initialized (tasks will respond immediately to stop signal)")
        
        # Создаем приложение
        self.create_application()
        
        # Запускаем периодический сбор в фоне (только в prod)
        collection_task = None
        tier_adjustment_task = None
        hourly_digest_task = None
        morning_digest_task = None
        from config.config import APP_ENV
        if APP_ENV == "prod":
            collection_task = asyncio.create_task(self.run_periodic_collection())
            tier_adjustment_task = asyncio.create_task(self.run_tier_adjustment())
            hourly_digest_task = asyncio.create_task(self.run_hourly_digest())
            morning_digest_task = asyncio.create_task(self.run_morning_digest())
        mgmt_runner = None
        
        # Запускаем приложение
        await self.application.initialize()
        await self.application.start()

        try:
            from config.railway_config import TG_MODE, WEBHOOK_BASE_URL, WEBHOOK_PATH, WEBHOOK_SECRET, PORT
        except (ImportError, ValueError):
            from config.config import TG_MODE, WEBHOOK_BASE_URL, WEBHOOK_PATH, WEBHOOK_SECRET, PORT

        tg_mode = (TG_MODE or "").strip().lower()
        if tg_mode in ("", "auto", "autodetect"):
            tg_mode = "webhook" if WEBHOOK_BASE_URL else "polling"
        elif tg_mode not in {"webhook", "polling"}:
            logger.warning(f"Unknown TG_MODE '{TG_MODE}', auto-detecting mode")
            tg_mode = "webhook" if WEBHOOK_BASE_URL else "polling"

        if tg_mode == "webhook":
            if not WEBHOOK_BASE_URL:
                logger.warning("WEBHOOK_BASE_URL missing for webhook mode, falling back to polling")
                tg_mode = "polling"
            else:
                webhook_url = WEBHOOK_BASE_URL.rstrip('/') + WEBHOOK_PATH
                await self.application.updater.start_webhook(
                    listen="0.0.0.0",
                    port=PORT,
                    url_path=WEBHOOK_PATH.lstrip('/'),
                    webhook_url=webhook_url,
                    secret_token=WEBHOOK_SECRET,
                )
                logger.info(f"Bot started with webhook: {webhook_url}")

        if tg_mode == "polling":
            try:
                await self.application.bot.delete_webhook(drop_pending_updates=False)
                logger.info("Deleted webhook before polling")
            except Exception as e:
                logger.warning(f"Failed to delete webhook before polling: {e}")
            await self.application.updater.start_polling()
            logger.info("Bot started with polling")

        try:
            from config.railway_config import MGMT_BIND, MGMT_PORT
        except (ImportError, ValueError):
            from config.config import MGMT_BIND, MGMT_PORT
        if MGMT_PORT:
            mgmt_runner = await start_mgmt_api(MGMT_BIND, MGMT_PORT)
            logger.info(f"Mgmt API started on {MGMT_BIND}:{MGMT_PORT}")
        
        try:
            await asyncio.Event().wait()  # Ждем завершения
        except KeyboardInterrupt:
            logger.info("Received interrupt signal")
        finally:
            self.is_running = False
            if collection_task:
                collection_task.cancel()
            await stop_mgmt_api(mgmt_runner)
            await self.application.updater.stop()
            await self.application.stop()
            await self.application.shutdown()
            self.db.release_bot_lock(self._db_instance_id)
            self._release_instance_lock()

    async def _shutdown_due_to_conflict(self, reason: str):
        """Shutdown bot immediately on 409 Conflict (duplicate instance)."""
        if self._shutdown_requested:
            return
        self._shutdown_requested = True
        logger.error(f"Shutting down due to конфликт: {reason}")
        try:
            self.is_running = False
            if self.application and self.application.updater:
                await self.application.updater.stop()
            if self.application:
                await self.application.stop()
                await self.application.shutdown()
        except Exception as e:
            logger.debug(f"Error during конфликт shutdown: {e}")
        finally:
            self.db.release_bot_lock(self._db_instance_id)
            self._release_instance_lock()
            os._exit(0)

    async def on_error(self, update: object, context: ContextTypes.DEFAULT_TYPE):
        """Global error handler for the bot."""
        err = getattr(context, "error", None)
        if isinstance(err, Conflict) or (err and "Conflict: terminated by other getUpdates request" in str(err)):
            await self._shutdown_due_to_conflict(str(err))
    async def _generate_doc_file(self, user_id: int) -> str | None:
        """
        Generate DOC file with selected news for user.
        Format: Title -> URL -> Tag -> Text (clean, minimal formatting)
        
        Args:
            user_id: Telegram user ID
            
        Returns:
            Path to generated file or None if error
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import tempfile
            
            selected_ids = self.db.get_user_selections(user_id, env="prod")
            if not selected_ids:
                return None
            
            # Create document with normal style throughout
            doc = Document()
            
            # Add header with generation date
            from datetime import datetime
            header_para = doc.add_paragraph(f"Создано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Set default font for entire document (Times New Roman, 12pt)
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Add each news
            for idx, news_id in enumerate(selected_ids):
                # Get news from DB or cache
                news = self.db.get_news_by_id(news_id) or self.news_cache.get(news_id)
                if not news:
                    continue
                
                # Add spacing between articles (not separator lines)
                if idx > 0:
                    doc.add_paragraph()
                
                # 1. Title
                title = news.get('title', 'Без заголовка').strip()
                title_para = doc.add_paragraph(title)
                for run in title_para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = False
                    run.font.color.rgb = None
                
                # 2. URL
                url = news.get('url', '').strip()
                if url:
                    url_para = doc.add_paragraph(url)
                    for run in url_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.color.rgb = None
                
                # 3. Tag (without emoji, just the hashtag)
                category = news.get('category', 'russia')
                category_map = {
                    'world': '#Мир',
                    'russia': '#Россия',
                    'moscow': '#Москва',
                    'moscow_region': '#Подмосковье',
                }
                tag = category_map.get(category, '#Россия')
                tag_para = doc.add_paragraph(tag)
                for run in tag_para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = False
                    run.font.color.rgb = None
                
                # 4. Text (AI summary if exists, otherwise original text)
                summary = self.db.get_cached_summary(news_id)
                text = summary if summary else news.get('clean_text') or news.get('text') or news.get('lead_text', 'Текст недоступен')
                text = text.strip()
                
                # Clean text: remove emoji and extra formatting
                import re
                text = re.sub(r'[😀-🙏🌀-🗿🚀-🛿]', '', text)  # Remove emoji
                text = re.sub(r'📰|🔗|💬|✉️|✅|❌|🤖|📄|📌|🌍|🇷🇺|🏛️|🏘️', '', text)  # Remove specific emoji
                text = re.sub(r'Источник:|Ссылка:|Тег:|Категория:|пересказ:|ИИ:|Оригинальный текст:', '', text, flags=re.IGNORECASE)  # Remove labels
                text = re.sub(r'\s+', ' ', text).strip()  # Clean up whitespace
                
                if text:
                    text_para = doc.add_paragraph(text)
                    for run in text_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = False
                        run.font.color.rgb = None
            
            # Save to temp file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Error generating DOC file: {e}", exc_info=True)
            return None

    async def _export_news_period(self, user_id: int, context: ContextTypes.DEFAULT_TYPE, hours: int) -> None:
        """Export news from the last N hours to Excel and send to user."""
        from datetime import datetime, timedelta

        try:
            end_dt = datetime.now()
            start_dt = end_dt - timedelta(hours=hours)

            news_items = self.db.get_news_in_period(start_dt, end_dt)
            if not news_items:
                await context.bot.send_message(
                    chat_id=user_id,
                    text=f"📭 За последние {hours} ч. новостей нет."
                )
                return

            excel_file = self._generate_excel_file_for_period(news_items)
            if not excel_file:
                await context.bot.send_message(chat_id=user_id, text="❌ Не удалось создать Excel")
                return

            filename = f"news_export_{hours}h_{end_dt.strftime('%Y%m%d_%H%M')}.xlsx"
            await context.bot.send_document(
                chat_id=user_id,
                document=open(excel_file, 'rb'),
                filename=filename,
                caption=f"📥 Unload: новости за последние {hours} ч. ({len(news_items)} шт.)"
            )

            import os
            os.remove(excel_file)

        except Exception as e:
            logger.error(f"Error exporting Excel: {e}")
            await context.bot.send_message(chat_id=user_id, text="❌ Ошибка при выгрузке")

    def _generate_excel_file_for_period(self, news_items: list) -> str | None:
        """Generate Excel file for news items list."""
        try:
            from utils.excel_export import generate_excel_file_for_period

            return generate_excel_file_for_period(news_items)
        except Exception as e:
            logger.error(f"Error generating Excel file: {e}")
            return None
    
    async def _show_sources_menu(self, query, page: int = 0):
        """Показать меню источников с пагинацией"""
        sources = self.db.list_sources()
        user_id = str(query.from_user.id)
        user_enabled = self.db.get_user_source_enabled_map(user_id, env="prod")
        
        # Пагинация
        PAGE_SIZE = 8
        total_pages = (len(sources) + PAGE_SIZE - 1) // PAGE_SIZE
        page = max(0, min(page, total_pages - 1))
        
        start = page * PAGE_SIZE
        end = start + PAGE_SIZE
        page_sources = sources[start:end]
        
        # Построить клавиатуру
        keyboard = []
        for src in page_sources:
            source_id = src['id']
            title = src['title']
            # Если нет записи в user_source_settings -> считаем True
            enabled = user_enabled.get(source_id, True)
            icon = "✅" if enabled else "⬜️"
            btn_text = f"{icon} {title}"
            keyboard.append([
                InlineKeyboardButton(btn_text, callback_data=f"settings:src_toggle:{source_id}:{page}")
            ])
        
        # Пагинация кнопок
        nav_buttons = []
        if page > 0:
            nav_buttons.append(InlineKeyboardButton("⬅️", callback_data=f"settings:src_page:{page-1}"))
        nav_buttons.append(InlineKeyboardButton(f"{page+1}/{total_pages}", callback_data="noop"))
        if page < total_pages - 1:
            nav_buttons.append(InlineKeyboardButton("➡️", callback_data=f"settings:src_page:{page+1}"))
        
        if nav_buttons:
            keyboard.append(nav_buttons)
        
        # Кнопка "Назад"
        keyboard.append([InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")])
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            text=f"📰 Источники новостей (страница {page+1}/{total_pages})\n\n✅ = включено\n⬜️ = отключено",
            reply_markup=reply_markup
        )
    
    def _filter_news_by_user_sources(self, news_items: list, user_id=None) -> list:
        """
        Отфильтровать новости по включённым для пользователя источникам.
        Если user_id=None или у пользователя все источники включены - возвращаем все.
        """
        if not user_id:
            return news_items
        
        enabled_source_ids = self.db.get_enabled_source_ids_for_user(user_id, env="prod")
        
        # Если None -> все включены
        if enabled_source_ids is None:
            return news_items
        
        # Преобразовать source_ids в set для быстрого поиска
        enabled_ids_set = set(enabled_source_ids)
        
        # Построить mapping source_code/title -> source_id
        sources = self.db.list_sources()
        code_to_id = {src['code']: src['id'] for src in sources}
        
        filtered = []
        for news in news_items:
            source = news.get('source', '')
            # Попробовать найти source_id по code или title
            source_id = code_to_id.get(source)
            if source_id and source_id in enabled_ids_set:
                filtered.append(news)
            elif not source_id:
                # Если источник не найден в БД - включаем его (по умолчанию)
                filtered.append(news)
        
        return filtered

    def _get_domain(self, news: dict) -> str:
        domain = news.get('domain')
        if domain:
            return str(domain).lower()
        url = news.get('url', '')
        try:
            return urlparse(url).netloc.lower() or "unknown"
        except Exception:
            return "unknown"

    def _record_drop_reason(self, domain: str, reason: str | None) -> None:
        if not reason:
            return
        bucket = self.drop_counters.setdefault(domain, {})
        bucket[reason] = bucket.get(reason, 0) + 1

    def _get_global_category_filter(self) -> str | None:
        return self.db.get_bot_setting("global_category_filter")

    def _set_global_category_filter(self, value: str | None) -> None:
        self.db.set_bot_setting("global_category_filter", value)

    def _get_user_category_filter(self, user_id: int) -> str | None:
        return self.db.get_user_category_filter(str(user_id), env="prod")

    def _set_user_category_filter(self, user_id: int, value: str | None) -> None:
        self.db.set_user_category_filter(str(user_id), value, env="prod")

    def _should_publish_news(self, news: dict) -> tuple[bool, str | None]:
        """Apply freshness rules based on published_confidence."""
        from datetime import datetime, timedelta

        domain = self._get_domain(news)
        freshness_days_override = {
            'new-science.ru': 7,
            'forklog.com': 3,
        }
        override_days = freshness_days_override.get(domain)
        confidence = (news.get('published_confidence') or 'none').lower()
        published_raw = news.get('published_at')
        published_at = parse_datetime_value(published_raw)
        published_date_raw = news.get('published_date')
        fallback_at = parse_datetime_value(news.get('fetched_at') or news.get('first_seen_at'))
        now_local = get_project_now()

        if not published_at and published_date_raw:
            try:
                pub_date = datetime.fromisoformat(str(published_date_raw)).date()
                published_at = datetime.combine(pub_date, datetime.min.time())
                confidence = 'low'
            except Exception:
                return False, "PARSE_DATE_FAILED"

        if not published_at:
            if published_raw:
                return False, "PARSE_DATE_FAILED"
            if fallback_at:
                published_at = fallback_at
                confidence = 'surrogate'
            else:
                return False, "NO_PUBLISHED_DATE"

        published_local = to_project_tz(published_at)
        if published_local > now_local + timedelta(minutes=5):
            return False, "PARSE_DATE_FAILED"

        if override_days is not None:
            if published_local < now_local - timedelta(days=override_days):
                return False, "OLD_PUBLISHED_AT"
            return True, None

        if published_local < now_local - timedelta(hours=36):
            return False, "OLD_PUBLISHED_AT"
        return True, None

    def _is_today_news(self, news: dict) -> bool:
        """Backward-compatible wrapper for freshness logic."""
        ok, _reason = self._should_publish_news(news)
        return ok

    async def _show_ai_management(self, query):
        """Show AI levels management screen"""
        try:
            from core.services.access_control import (
                get_effective_level,
                get_global_level,
                get_user_level_override,
            )

            app_env = get_app_env()
            user_id = str(query.from_user.id)

            if app_env == "sandbox":
                # Admin-only global management
                is_admin = self._is_admin(int(user_id))
                if not is_admin:
                    await query.answer("❌ Доступ запрещён", show_alert=True)
                    return

                hashtags_level = get_global_level(self.db, 'hashtags')
                cleanup_level = get_global_level(self.db, 'cleanup')
                summary_level = get_global_level(self.db, 'summary')
            else:
                hashtags_level = get_effective_level(self.db, user_id, 'hashtags')
                cleanup_level = get_effective_level(self.db, user_id, 'cleanup')
                summary_level = get_effective_level(self.db, user_id, 'summary')

                hashtags_global = get_global_level(self.db, 'hashtags')
                cleanup_global = get_global_level(self.db, 'cleanup')
                summary_global = get_global_level(self.db, 'summary')

                hashtags_override = get_user_level_override(self.db, user_id, 'hashtags')
                cleanup_override = get_user_level_override(self.db, user_id, 'cleanup')
                summary_override = get_user_level_override(self.db, user_id, 'summary')
            
            # Build UI
            def level_text(level: int) -> str:
                return "OFF" if level == 0 else str(level)
            
            def level_icon(level: int) -> str:
                return "⬜️" if level == 0 else "✅"
            
            keyboard = []
            
            # Hashtags
            header = f"{level_icon(hashtags_level)} 🏷 Хештеги (AI): {level_text(hashtags_level)}"
            if app_env == "prod":
                header += f" | G {level_text(hashtags_global)}"
                if hashtags_override is not None:
                    header += f" | U {level_text(hashtags_override)}"
            keyboard.append([InlineKeyboardButton(header, callback_data="noop")])
            keyboard.append([
                InlineKeyboardButton("−", callback_data="ai:dec:hashtags"),
                InlineKeyboardButton("OFF", callback_data="ai:set:hashtags:0"),
                InlineKeyboardButton("+", callback_data="ai:inc:hashtags"),
            ])
            
            # Cleanup
            header = f"{level_icon(cleanup_level)} 🧹 Очистка (AI): {level_text(cleanup_level)}"
            if app_env == "prod":
                header += f" | G {level_text(cleanup_global)}"
                if cleanup_override is not None:
                    header += f" | U {level_text(cleanup_override)}"
            keyboard.append([InlineKeyboardButton(header, callback_data="noop")])
            keyboard.append([
                InlineKeyboardButton("−", callback_data="ai:dec:cleanup"),
                InlineKeyboardButton("OFF", callback_data="ai:set:cleanup:0"),
                InlineKeyboardButton("+", callback_data="ai:inc:cleanup"),
            ])
            
            # Summary
            header = f"{level_icon(summary_level)} 📝 Пересказ (AI): {level_text(summary_level)}"
            if app_env == "prod":
                header += f" | G {level_text(summary_global)}"
                if summary_override is not None:
                    header += f" | U {level_text(summary_override)}"
            keyboard.append([InlineKeyboardButton(header, callback_data="noop")])
            keyboard.append([
                InlineKeyboardButton("−", callback_data="ai:dec:summary"),
                InlineKeyboardButton("OFF", callback_data="ai:set:summary:0"),
                InlineKeyboardButton("+", callback_data="ai:inc:summary"),
            ])
            
            # Back button
            keyboard.append([InlineKeyboardButton("⬅️ Назад", callback_data="settings:back")])
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            
            if app_env == "sandbox":
                text = (
                    "🤖 Управление AI модулями (ГЛОБАЛЬНЫЕ)\n\n"
                    "Уровни 0-5:\n"
                    "• 0 = выключено (no LLM calls)\n"
                    "• 1-2 = быстрый/экономный режим\n"
                    "• 3 = стандартный (по умолчанию)\n"
                    "• 4-5 = максимальное качество\n\n"
                    "⚡️ Очистка level=5: автоматический пересказ\n"
                    "   для lenta.ru и ria.ru (1-2 предложения)\n\n"
                    "Используйте − и + для настройки уровня,\n"
                    "или OFF для полного отключения.\n\n"
                    "⚠️ Настройки применяются к ПРОДУ и ПЕСОЧНИЦЕ"
                )
            else:
                text = (
                    "🤖 Управление AI модулями (ПЕРСОНАЛЬНЫЕ)\n\n"
                    "Уровни 0-5:\n"
                    "• 0 = выключено (no LLM calls)\n"
                    "• 1-2 = быстрый/экономный режим\n"
                    "• 3 = стандартный (по умолчанию)\n"
                    "• 4-5 = максимальное качество\n\n"
                    "G = global, U = user override\n"
                    "Персональные уровни действуют только для вас."
                )
            
            await query.edit_message_text(text=text, reply_markup=reply_markup)
        except Exception as e:
            logger.error(f"AI management error: {e}")
            await query.answer("❌ Ошибка меню AI", show_alert=True)
    
    async def _handle_ai_level_change(self, query, module: str, action: str, level: int = None):
        """Handle AI level change (inc/dec/set) - uses global settings"""
        from core.services.access_control import (
            get_effective_level,
            set_global_level,
            set_user_level,
        )

        user_id = str(query.from_user.id)
        app_env = get_app_env()

        if app_env == "sandbox":
            is_admin = self._is_admin(int(user_id))
            if not is_admin:
                await query.answer("❌ Доступ запрещён", show_alert=True)
                return

            if action == "inc":
                current = get_effective_level(self.db, 'global', module)
                new_level = min(5, current + 1)
                set_global_level(self.db, module, new_level)
            elif action == "dec":
                current = get_effective_level(self.db, 'global', module)
                new_level = max(0, current - 1)
                set_global_level(self.db, module, new_level)
            elif action == "set":
                new_level = max(0, min(5, int(level)))
                set_global_level(self.db, module, new_level)
            else:
                await query.answer("❌ Invalid action", show_alert=True)
                return
        else:
            if action == "inc":
                current = get_effective_level(self.db, user_id, module)
                new_level = min(5, current + 1)
                set_user_level(self.db, user_id, module, new_level)
            elif action == "dec":
                current = get_effective_level(self.db, user_id, module)
                new_level = max(0, current - 1)
                set_user_level(self.db, user_id, module, new_level)
            elif action == "set":
                new_level = max(0, min(5, int(level)))
                set_user_level(self.db, user_id, module, new_level)
            else:
                await query.answer("❌ Invalid action", show_alert=True)
                return

        await query.answer(f"✅ {module}: {new_level}")
        await self._show_ai_management(query)

    async def _finalize_invite_creation(
        self,
        admin_id: str,
        label: str | None,
        context: ContextTypes.DEFAULT_TYPE,
        update: Update | None = None,
        query=None,
    ) -> None:
        cleaned_label = label.strip() if label else None
        if cleaned_label and len(cleaned_label) > 80:
            cleaned_label = cleaned_label[:80]

        invite_code = self._generate_signed_invite_code(admin_id)
        if not invite_code:
            message = (
                "❌ Не задан INVITE_SECRET.\n\n"
                "Установите переменную окружения INVITE_SECRET одинаково в prod и sandbox."
            )
            back_markup = InlineKeyboardMarkup([
                [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")]
            ])
            if query:
                await query.edit_message_text(message, reply_markup=back_markup)
            elif update:
                await update.message.reply_text(message, reply_markup=back_markup)
            return

        self.db.create_invite_with_code(invite_code, admin_id, invite_label=cleaned_label)
        try:
            self.access_db.create_invite_with_code(invite_code, admin_id, invite_label=cleaned_label)
        except Exception:
            pass

        # Get bot username for link
        try:
            from config.railway_config import BOT_PROD_USERNAME
        except (ImportError, ValueError):
            try:
                from config.config import BOT_PROD_USERNAME
            except ImportError:
                BOT_PROD_USERNAME = None

        if not BOT_PROD_USERNAME:
            bot_info = await self.application.bot.get_me()
            bot_username = bot_info.username
        else:
            bot_username = BOT_PROD_USERNAME

        invite_link = f"https://t.me/{bot_username}?start={invite_code}"
        if cleaned_label:
            from html import escape
            label_line = f"👤 Для: {escape(cleaned_label)}\n"
        else:
            label_line = ""

        keyboard = [
            [InlineKeyboardButton("📤 Отправить пользователю", callback_data=f"mgmt:send_invite:{invite_code}")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:users")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)

        message = (
            "🎉 Новый инвайт-код создан!\n\n"
            f"{label_line}"
            f"📌 Код: <code>{invite_code}</code>\n\n"
            "🔗 Ссылка для пользователя:\n"
            f"<code>{invite_link}</code>\n\n"
            "Нажмите кнопку ниже для отправки инвайта пользователю."
        )

        if query:
            await query.edit_message_text(text=message, reply_markup=reply_markup, parse_mode='HTML')
        elif update:
            await update.message.reply_text(message, reply_markup=reply_markup, parse_mode='HTML')

    async def _show_users_management(self, query):
        """Show users and invites management screen"""
        try:
            from config.railway_config import APP_ENV
        except (ImportError, ValueError):
            from config.config import APP_ENV

        user_id = query.from_user.id

        # Check admin
        is_admin = self._is_admin(user_id)
        if not is_admin:
            await query.answer("❌ Доступ запрещён", show_alert=True)
            return

        # For prod, sandbox restriction should not apply (admins can manage both)
        # Get invites and approved users from DB
        unused_invites = self.db.get_unused_invites()
        approved_users = self.access_db.get_approved_users()

        # Build UI
        keyboard = []

        # Users section
        keyboard.append([InlineKeyboardButton("👥 Одобренные пользователи", callback_data="mgmt:users_list")])
        if approved_users:
            keyboard.append([InlineKeyboardButton(f"({len(approved_users)} чел.)", callback_data="noop")])
        else:
            keyboard.append([InlineKeyboardButton("(нет)", callback_data="noop")])

        # Invites section
        keyboard.append([InlineKeyboardButton("📨 Активные инвайты", callback_data="noop")])
        if unused_invites:
            keyboard.append([InlineKeyboardButton(f"({len(unused_invites)} инвайтов)", callback_data="noop")])
        else:
            keyboard.append([InlineKeyboardButton("(нет)", callback_data="noop")])

        # Action buttons
        keyboard.append([
            InlineKeyboardButton("➕ Создать инвайт", callback_data="mgmt:new_invite"),
            InlineKeyboardButton("👁️ Список", callback_data="mgmt:users_list"),
        ])

        # Back button
        keyboard.append([InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:back")])

        reply_markup = InlineKeyboardMarkup(keyboard)

        text = (
            "👥 Управление пользователями и инвайтами\n\n"
            f"✅ Одобренные: {len(approved_users)} чел.\n"
            f"📨 Активные инвайты: {len(unused_invites)}\n\n"
            "Используйте кнопки ниже для управления."
        )

        await query.edit_message_text(text=text, reply_markup=reply_markup)

    async def cmd_management_inline(self, query):
        """Show main management menu via inline query"""
        reply_markup = self._build_sandbox_admin_keyboard()
        await query.edit_message_text(
            text="🛠 Управление системой",
            reply_markup=reply_markup
        )
    
    async def _trigger_news_collection(self):
        """Запускает сбор новостей в песочнице (фоновая задача)"""
        try:
            logger.info("[ADMIN] Triggering news collection in sandbox...")
            
            # Запускаем сбор с force=True чтобы игнорировать таймауты
            if self.collector:
                results = await self.collector.collect_all(force=True)
                logger.info(f"[ADMIN] Collection completed: {len(results)} items collected")
            else:
                logger.warning("[ADMIN] Collector not available")
        except Exception as e:
            logger.error(f"[ADMIN] Error during news collection: {e}", exc_info=True)

    async def _show_admin_status(self, query):
        """📊 System status panel"""
        from core.services.global_stop import get_global_stop_status_str, is_redis_available
        
        try:
            app_env = get_app_env()
            is_stopped, stop_status = get_global_stop_status_str()
            redis_ok = is_redis_available()
            
            # Build status text with emojis
            status_lines = [
                "📊 СТАТУС СИСТЕМЫ",
                "",
                f"🤖 Окружение: {app_env.upper()}",
                f"⏹ Глобальная остановка: {stop_status}",
                f"� Redis статус: {'✅ OK' if redis_ok else '⚠️ Fallback (SQLite)'}",
                f"🗄️ БД: SQLite (news.db)",
                "",
                "Нажмите кнопки ниже для управления системой."
            ]
            
            text = "\n".join(status_lines)
            
            # Build keyboard with toggle global stop button
            keyboard = [
                [InlineKeyboardButton(
                    "⛔ ОСТАНОВИТЬ ВСЮ СИСТЕМУ" if not is_stopped else "✅ ВОЗОБНОВИТЬ ВСЮ СИСТЕМУ",
                    callback_data="mgmt:toggle_global_stop",
                )],
                [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")],
            ]
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=text, reply_markup=reply_markup)
            
        except Exception as e:
            logger.error(f"Error in _show_admin_status: {e}")
            await query.answer("❌ Ошибка при загрузке статуса", show_alert=True)

    async def _show_admin_ai_panel(self, query):
        """🤖 AI management panel"""
        text = (
            "🤖 УПРАВЛЕНИЕ AI МОДУЛЯМИ\n\n"
            "Доступные модули:\n"
            "• Hashtags - автоматическое добавление хешTagов\n"
            "• Cleanup - очистка текста\n"
            "• Summary - генерация резюме\n\n"
            "Уровень: 0 = отключено, 5 = максимум\n\n"
            "Выберите модуль для управления:"
        )
        
        keyboard = [
            [InlineKeyboardButton("🏷️ Hashtags", callback_data="mgmt:ai:module:hashtags")],
            [InlineKeyboardButton("🧹 Cleanup", callback_data="mgmt:ai:module:cleanup")],
            [InlineKeyboardButton("📝 Summary", callback_data="mgmt:ai:module:summary")],
            [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")],
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup)

    async def _show_admin_sources_panel(self, query, page: int = 0):
        """📰 Sources management panel with paginated list"""
        sources = self.db.list_sources()
        if not sources:
            await query.edit_message_text(
                text="📰 Нет источников в базе данных",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")]
                ])
            )
            return
        
        # Sort by quality_score DESC
        sources_with_quality = []
        for src in sources:
            quality = self.db.get_source_quality(src['code']) or {}
            quality_score = quality.get('quality_score', 0.0)
            error_streak = quality.get('error_streak', 0)
            last_success = quality.get('last_success_at', 'никогда')
            tier = self.db.get_source_tier(src['code'])
            sources_with_quality.append({
                **src,
                'quality_score': quality_score,
                'error_streak': error_streak,
                'last_success': last_success,
                'tier': tier
            })
        
        sources_with_quality.sort(key=lambda x: x['quality_score'], reverse=True)
        
        # Pagination
        per_page = 8
        total_pages = (len(sources_with_quality) + per_page - 1) // per_page
        page = max(0, min(page, total_pages - 1))
        start_idx = page * per_page
        end_idx = min(start_idx + per_page, len(sources_with_quality))
        page_sources = sources_with_quality[start_idx:end_idx]
        
        # Build message
        text_lines = [
            f"📰 ИСТОЧНИКИ ({len(sources)} всего)\n",
            f"Страница {page + 1}/{total_pages}\n"
        ]
        
        keyboard = []
        
        for src in page_sources:
            # Status indicator
            if src.get('enabled_global'):
                if src['error_streak'] >= 5:
                    status = "🔴"  # Enabled but failing
                elif src['error_streak'] > 0:
                    status = "🟡"  # Enabled with some errors
                else:
                    status = "🟢"  # OK
            else:
                status = "⚫"  # Disabled
            
            # Format source line
            name = src['name'][:20] if len(src['name']) > 20 else src['name']
            tier = src.get('tier', 'B')
            score = src.get('quality_score', 0.0)
            
            button_text = f"{status} {name} [{'A' if tier=='A' else 'B' if tier=='B' else 'C'}] {score:.2f}"
            keyboard.append([InlineKeyboardButton(
                button_text,
                callback_data=f"mgmt:source:detail:{src['code']}"
            )])
        
        # Pagination buttons
        nav_buttons = []
        if page > 0:
            nav_buttons.append(InlineKeyboardButton("◀️", callback_data=f"mgmt:sources:page:{page-1}"))
        nav_buttons.append(InlineKeyboardButton(f"{page+1}/{total_pages}", callback_data="noop"))
        if page < total_pages - 1:
            nav_buttons.append(InlineKeyboardButton("▶️", callback_data=f"mgmt:sources:page:{page+1}"))
        keyboard.append(nav_buttons)
        
        # Action buttons
        keyboard.append([
            InlineKeyboardButton("🔍 Тест всех", callback_data="mgmt:sources:test_all"),
            InlineKeyboardButton("🔄 Авто-тюнинг", callback_data="mgmt:sources:auto_tune")
        ])
        keyboard.append([InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")])
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        text = "\n".join(text_lines) + "\n\nЛегенда:\n🟢 OK | 🟡 Ошибки | 🔴 Много ошибок | ⚫ Выключен\n[A/B/C] = tier"
        await query.edit_message_text(text=text, reply_markup=reply_markup)

    async def _show_source_detail(self, query, source_code: str):
        """Show detailed information about a source with action buttons."""
        # Get source info
        sources = self.db.list_sources()
        source = next((s for s in sources if s['code'] == source_code), None)
        if not source:
            await query.answer("❌ Источник не найден", show_alert=True)
            return
        
        # Check if quarantined
        is_quarantined = self.db.is_source_quarantined(source_code)
        quarantine_at = source.get('quarantined_at')
        quarantine_reason = source.get('quarantine_reason', 'unknown')
        
        # Get quality metrics
        quality = self.db.get_source_quality(source_code) or {}
        quality_score = quality.get('quality_score', 0.0)
        success_count = quality.get('success_count', 0)
        error_count = quality.get('error_count', 0)
        error_streak = quality.get('error_streak', 0)
        items_total = quality.get('items_total', 0)
        items_new = quality.get('items_new', 0)
        items_duplicate = quality.get('items_duplicate', 0)
        last_success = quality.get('last_success_at', 'никогда')
        last_error = quality.get('last_error_at', 'нет')
        last_error_code = quality.get('last_error_code', '-')
        
        # Get tier
        tier = self.db.get_source_tier(source_code)
        tier_params = self.db.get_tier_params(tier)
        
        # Status indicator
        if is_quarantined:
            status = "🔴 В КАРАНТИНЕ"
        elif source.get('enabled_global'):
            if error_streak >= 5:
                status = "🔴 Активен (много ошибок)"
            elif error_streak > 0:
                status = "🟡 Активен (есть ошибки)"
            else:
                status = "🟢 Активен"
        else:
            status = "⚫ Отключен"
        
        # Format message
        text = (
            f"📰 {source['name']}\n\n"
            f"Код: {source_code}\n"
            f"Статус: {status}\n"
            f"Tier: {tier} (интервал: {tier_params['min_interval_seconds']//60} мин, "
            f"лимит: {tier_params.get('max_items_per_fetch', 'нет')})\n\n"
        )
        
        # Add quarantine info if quarantined
        if is_quarantined:
            text += (
                f"⚠️ КАРАНТИН:\n"
                f"   Причина: {quarantine_reason}\n"
                f"   С: {quarantine_at or 'N/A'}\n\n"
            )
        
        text += (
            f"📊 Качество: {quality_score:.3f}\n"
            f"✅ Успехов: {success_count}\n"
            f"❌ Ошибок: {error_count} (streak: {error_streak})\n\n"
            f"📈 Статистика:\n"
            f"  Всего получено: {items_total}\n"
            f"  Новых: {items_new}\n"
            f"  Дубликатов: {items_duplicate}\n\n"
            f"🕐 Последний успех: {last_success}\n"
            f"❗ Последняя ошибка: {last_error}\n"
            f"   Код ошибки: {last_error_code}"
        )
        
        # Build keyboard
        keyboard = []
        
        # If quarantined, show restore button
        if is_quarantined:
            keyboard.append([InlineKeyboardButton("🔓 Восстановить", callback_data=f"mgmt:source:restore:{source_code}")])
        else:
            # Toggle enable/disable
            toggle_text = "⛔ Отключить" if source.get('enabled_global') else "✅ Включить"
            keyboard.append([InlineKeyboardButton(toggle_text, callback_data=f"mgmt:source:toggle:{source_code}")])
        
        # Tier change buttons
        tier_buttons = []
        for t in ['A', 'B', 'C']:
            if t == tier:
                tier_buttons.append(InlineKeyboardButton(f"[{t}]", callback_data="noop"))
            else:
                tier_buttons.append(InlineKeyboardButton(f"{t}", callback_data=f"mgmt:source:tier:{source_code}:{t}"))
        keyboard.append(tier_buttons)
        
        # Test fetch button
        keyboard.append([InlineKeyboardButton("🧪 Тест сбора", callback_data=f"mgmt:source:test:{source_code}")])
        
        # Back button
        keyboard.append([InlineKeyboardButton("⬅️ К списку", callback_data="mgmt:sources")])
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup)

    async def _show_admin_stats_panel(self, query):
        """📈 Statistics panel"""
        # Get database stats
        db_stats = self.db.get_stats()
        ai_stats = self.db.get_ai_usage()
        
        # Get top source
        try:
            cursor = self.db._conn.cursor()
            cursor.execute('''
                SELECT source, COUNT(*) as cnt 
                FROM published_news 
                WHERE published_at > datetime('now', '-1 day')
                GROUP BY source 
                ORDER BY cnt DESC 
                LIMIT 1
            ''')
            top_source_row = cursor.fetchone()
            top_source = f"{top_source_row[0]} ({top_source_row[1]})" if top_source_row else "-"
        except Exception as e:
            logger.error(f"Error getting top source: {e}")
            top_source = "-"
        
        text = (
            "📈 СТАТИСТИКА\n\n"
            f"📊 Опубликовано (24ч): {db_stats.get('today', 0)} новостей\n"
            f"🤖 Использование AI: {ai_stats['total_requests']} запросов\n"
            f"💰 Расходы AI: ${ai_stats.get('total_cost_usd', 0):.4f}\n"
            f"⚠ Ошибок/предупреждений: 0\n"
            f"📰 Топ источник: {top_source}\n\n"
            "Обновляется каждые 5 минут."
        )
        
        keyboard = [
            [InlineKeyboardButton("🔄 Обновить", callback_data="mgmt:stats:refresh")],
            [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")],
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup)

    async def _show_admin_diagnostics_panel(self, query):
        """🧰 Diagnostics panel"""
        from core.services.global_stop import is_redis_available, get_global_stop
        try:
            try:
                from config.railway_config import RSSHUB_BASE_URL
            except (ImportError, ValueError):
                from config.config import RSSHUB_BASE_URL

            redis_ok = is_redis_available()
            global_stop = get_global_stop()
            rsshub_url = RSSHUB_BASE_URL or "-"

            text = (
                "🧰 ДИАГНОСТИКА\n\n"
                f"� Redis: {'✅ OK' if redis_ok else '⚠️ Fallback (SQLite)'}\n"
                f"⛔ Global stop: {'ON' if global_stop else 'OFF'}\n"
                f"🗄️ DB: {self.db.db_path}\n"
                f"🛰 RSSHub: {rsshub_url}\n"
            )

            keyboard = [
                [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")],
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.edit_message_text(text=text, reply_markup=reply_markup)
        except Exception as e:
            logger.error(f"Diagnostics panel error: {e}")
            await query.answer("❌ Ошибка диагностики", show_alert=True)

    async def _show_admin_settings_panel(self, query):
        """⚙ Settings panel"""
        try:
            from config.railway_config import CHECK_INTERVAL_SECONDS
        except (ImportError, ValueError):
            from config.config import CHECK_INTERVAL_SECONDS
        
        text = (
            "⚙ НАСТРОЙКИ СИСТЕМЫ\n\n"
            f"⏱️ Интервал проверки: {CHECK_INTERVAL_SECONDS}с\n"
            f"🔄 Параллельные задачи: 3\n"
            f"📝 Уровень логирования: INFO\n\n"
            "Выберите параметр для изменения:"
        )
        
        keyboard = [
            [InlineKeyboardButton("⏱️ Интервал проверки", callback_data="mgmt:settings:interval")],
            [InlineKeyboardButton("🔄 Параллельность", callback_data="mgmt:settings:parallel")],
            [InlineKeyboardButton("📝 Логирование", callback_data="mgmt:settings:logging")],
            [InlineKeyboardButton("⬅️ Назад в меню", callback_data="mgmt:main")],
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup)
    async def _show_ai_module_control(self, query, module: str):
        """Show control panel for specific AI module"""
        # Get current level from AILevelManager
        from core.services.access_control import AILevelManager
        ai_manager = AILevelManager(self.db)
        current_level = ai_manager.get_level('global', module)
        
        text = (
            f"🤖 МОДУЛЬ: {module.upper()}\n\n"
            f"Текущий уровень: {current_level} из 5\n"
            f"0 = отключено\n"
            f"5 = максимальный\n\n"
            "Выберите новый уровень:"
        )
        
        keyboard = [
            [
                InlineKeyboardButton("0", callback_data=f"mgmt:ai:level:{module}:set:0"),
                InlineKeyboardButton("1", callback_data=f"mgmt:ai:level:{module}:set:1"),
                InlineKeyboardButton("2", callback_data=f"mgmt:ai:level:{module}:set:2"),
                InlineKeyboardButton("3", callback_data=f"mgmt:ai:level:{module}:set:3"),
            ],
            [
                InlineKeyboardButton("4", callback_data=f"mgmt:ai:level:{module}:set:4"),
                InlineKeyboardButton("5", callback_data=f"mgmt:ai:level:{module}:set:5"),
            ],
            [InlineKeyboardButton("⬅️ Назад", callback_data="mgmt:ai")],
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(text=text, reply_markup=reply_markup)